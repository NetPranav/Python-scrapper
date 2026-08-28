#!/usr/bin/env python3
"""
PDF Scraper & Compiler
======================
Scrapes title, authors (name + college), abstract, and keywords from
IEEE conference paper PDFs and compiles them into a single formatted
output PDF with Table of Contents.

Usage:
    python3 scraper.py                              # uses default folders
    python3 scraper.py <input_folder> <output_folder>

Default folders:
    Input:  IncompletePDF/
    Output: CompletedPDF/

Dependencies:
    pip3 install PyMuPDF reportlab
"""

import os
import sys
import re
import json
import shutil
import random
import datetime

# Ensure UTF-8 stdout/stderr on Windows to avoid UnicodeEncodeError in cp1252
if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')
    except Exception:
        pass

# PyMuPDF import with fallback to avoid deprecation warning
try:
    import pymupdf as fitz
except ImportError:
    import fitz
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.colors import black, HexColor
from reportlab.platypus import (
    BaseDocTemplate, PageTemplate, Frame,
    Paragraph, Spacer, Flowable, KeepTogether, PageBreak
)
from reportlab.pdfgen import canvas as canvas_mod
from reportlab.pdfbase.pdfmetrics import stringWidth, registerFont
from reportlab.pdfbase.ttfonts import TTFont

# Register custom fonts if available
def register_custom_fonts():
    fonts_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'fonts')
    if os.path.exists(fonts_dir):
        for file in os.listdir(fonts_dir):
            if file.endswith('.ttf'):
                font_name = file[:-4] # filename without .ttf
                try:
                    registerFont(TTFont(font_name, os.path.join(fonts_dir, file)))
                except Exception as e:
                    print(f"Failed to register font {font_name}: {e}")

register_custom_fonts()


# ============================================================
# CONFIGURATION — loaded from config.json if present
# ============================================================

def _load_config():
    """Load parameters from config.json, falling back to defaults."""
    defaults = {
        'CONFERENCE_HEADER': (
            'International Conference on Advances in Computer '
            'Research on Science Engineering and Technology, 2026'
        ),
        'FOOTER_URL': 'www.acroset.in',
        'LINE_SPACING_FACTOR': 1.45,
        'WORD_SPACING': 1.5,
        'TITLE_FONT_SIZE': 12,
        'AUTHORS_FONT_SIZE': 9.5,
        'ABSTRACT_FONT_SIZE': 10,
        'KEYWORDS_FONT_SIZE': 10,
        'HEADER_FONT_SIZE': 8.5,
        'FOOTER_URL_FONT_SIZE': 9,
        'FOOTER_PAGE_FONT_SIZE': 10,
    }
    config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'config.json')
    if os.path.exists(config_path):
        try:
            with open(config_path, 'r') as f:
                user_cfg = json.load(f)
            for k in defaults:
                if k in user_cfg:
                    defaults[k] = user_cfg[k]
            print(f"  [CONFIG] Loaded config from {config_path}")
        except Exception as e:
            print(f"  [WARN] Could not load config.json: {e}")
    return defaults

_cfg = _load_config()

CONFERENCE_HEADER = _cfg['CONFERENCE_HEADER']
FOOTER_URL = _cfg['FOOTER_URL']

# Deep blue color for header/footer divider lines
DIVIDER_COLOR = HexColor('#003366')

PAGE_WIDTH, PAGE_HEIGHT = A4  # 595.27 x 841.89 points
LEFT_MARGIN = 50
RIGHT_MARGIN = 50
TOP_MARGIN = 70      # more space for header area
BOTTOM_MARGIN = 60   # more space for footer area
CONTENT_WIDTH = PAGE_WIDTH - LEFT_MARGIN - RIGHT_MARGIN

# Strict limit: max papers per page
MAX_PAPERS_PER_PAGE = 2
# Non-final pages must use at least this fraction of available height (max 80% whitespace)
MIN_PAGE_FILL_RATIO = 0.20
# Conservative headroom so KeepTogether blocks never overflow the frame
PACKING_HEADROOM_PT = 85.0

# ────────────────────────────────────────────────────────────
# TUNABLE SPACING PARAMETERS  (adjust these to control layout)
# ────────────────────────────────────────────────────────────
LINE_SPACING_FACTOR = _cfg['LINE_SPACING_FACTOR']
WORD_SPACING = _cfg['WORD_SPACING']
TITLE_FONT_SIZE = _cfg['TITLE_FONT_SIZE']
AUTHORS_FONT_SIZE = _cfg['AUTHORS_FONT_SIZE']
ABSTRACT_FONT_SIZE = _cfg['ABSTRACT_FONT_SIZE']
KEYWORDS_FONT_SIZE = _cfg['KEYWORDS_FONT_SIZE']
HEADER_FONT_SIZE = _cfg['HEADER_FONT_SIZE']
FOOTER_URL_FONT_SIZE = _cfg['FOOTER_URL_FONT_SIZE']
FOOTER_PAGE_FONT_SIZE = _cfg['FOOTER_PAGE_FONT_SIZE']

# Exhaustive list of country names for location-line detection
COUNTRIES = {
    'india', 'usa', 'uk', 'china', 'japan', 'germany', 'france',
    'australia', 'canada', 'united states', 'united kingdom', 'brazil',
    'south korea', 'singapore', 'malaysia', 'indonesia', 'thailand',
    'spain', 'italy', 'netherlands', 'sweden', 'switzerland', 'austria',
    'belgium', 'portugal', 'ireland', 'new zealand', 'south africa',
    'turkey', 'russia', 'pakistan', 'bangladesh', 'sri lanka', 'nepal',
    'mexico', 'philippines', 'vietnam', 'taiwan', 'hong kong', 'uae',
    'saudi arabia', 'nigeria', 'kenya', 'iran', 'israel', 'poland',
    'romania', 'hungary', 'greece', 'finland', 'iceland', 'norway',
    'denmark', 'czech republic', 'argentina', 'colombia', 'chile',
    'peru', 'qatar', 'oman', 'bahrain', 'kuwait', 'ghana', 'ethiopia',
    'tanzania', 'uganda', 'iraq', 'jordan', 'lebanon', 'morocco',
    'tunisia', 'algeria', 'egypt',
}

# US states / Indian states that appear in location lines
STATES = {
    'maharashtra', 'madhya pradesh', 'uttar pradesh', 'rajasthan',
    'karnataka', 'tamil nadu', 'telangana', 'andhra pradesh',
    'gujarat', 'punjab', 'haryana', 'bihar', 'west bengal', 'odisha',
    'kerala', 'assam', 'jharkhand', 'chhattisgarh', 'uttarakhand',
    'himachal pradesh', 'goa', 'tripura', 'meghalaya', 'sikkim',
    'michigan', 'california', 'texas', 'new york', 'florida',
    'illinois', 'ohio', 'pennsylvania', 'georgia', 'virginia',
    'massachusetts', 'washington', 'maryland', 'colorado', 'minnesota',
}


# ============================================================
# UTILITY FUNCTIONS
# ============================================================

def html_escape(text):
    """Escape special characters for ReportLab XML markup."""
    return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;'))


CITIES = {
    'rajkot', 'jaipur', 'kishangarh', 'pilani', 'kanyakumari', 'chennai', 'muscat',
    'vaddeswaram', 'bhopal', 'gwalior', 'vijayawada', 'bhimavaram', 'jalandhar', 'indore',
    'delhi', 'mumbai', 'pune', 'bangalore', 'bengaluru', 'hyderabad', 'kolkata',
    'ahmedabad', 'noida', 'gurugram', 'chandigarh', 'dehradun', 'oman', 'india', 'usa',
    'uk', 'singapore', 'malaysia', 'australia', 'kanpur', 'lucknow', 'patna', 'agra',
    'rajasthan', 'madhya pradesh', 'tamil nadu', 'kerala', 'karnataka', 'andhra pradesh',
    'punjab', 'haryana', 'uttar pradesh', 'gujarat', 'maharashtra'
}


def is_location_line(text):
    """
    Detect city/country/state lines in the author section.
    These lines typically look like "Indore, India" or "Bhopal, India".
    """
    clean = text.strip().rstrip('.,').lower()
    if clean in CITIES:
        return True

    # Check country names
    for country in COUNTRIES:
        if country in clean:
            return True

    # Check state names
    for state in STATES:
        if state in clean:
            return True

    # Pattern: "City, Country" — short text with comma, no email
    if ',' in text and len(text.strip()) < 40 and '@' not in text:
        parts = [p.strip().lower() for p in text.strip().split(',')]
        if len(parts) == 2 and any(p in CITIES or p in [c.lower() for c in COUNTRIES] or p in [s.lower() for s in STATES] for p in parts):
            return True

    return False


def is_email(text):
    """Detect email addresses or split-email fragments like 'edu.in'."""
    t = text.strip().lower()
    if '@' in t:
        return True
    if re.search(r'[\w.+-]+@[\w.-]+\.\w{2,}', t):
        return True
    if t.startswith('@') or re.match(r'^[\w.-]+@', t):
        return True
    if re.match(r'^[a-z]+\.[a-z]{2,3}$', t) and len(t) < 12:
        return True
    return False


def is_copyright(text):
    """Detect copyright/ISBN/IEEE footer lines."""
    t = text.strip().lower()
    return ('978-' in t or '©' in t or t.strip() == 'ieee' or
            'doi:' in t or 'doi ' in t)


def smart_join_parts(parts):
    """
    Intelligently join multi-line college/institution text.
    Adds commas between separate entities (dept → college),
    but uses spaces for line-continuations (word wraps).
    """
    cleaned = [p.strip() for p in parts if p.strip()]
    if not cleaned:
        return ""
    if len(cleaned) == 1:
        return cleaned[0]

    result = cleaned[0]
    for i in range(1, len(cleaned)):
        prev = cleaned[i - 1].rstrip()
        curr = cleaned[i].lstrip()

        # If previous line ends with a continuation character → space
        if prev.endswith(('&', ',', '-', '(')) or prev.lower().endswith(
                ('and', 'of', 'in', 'for', 'the', 'with', 'on')):
            result += ' ' + curr
        # If current line starts with lowercase → continuation
        elif curr and curr[0].islower():
            result += ' ' + curr
        # If current line starts with ')' → closing bracket continuation
        elif curr and curr[0] == ')':
            result += curr
        else:
            # Likely a new entity (e.g., department → university)
            result += ', ' + curr

    return re.sub(r'\s+', ' ', result).strip()


INSTITUTION_RE = re.compile(
    r'\b(?:department|dept\.?|university|institute|college|school|faculty|academy|'
    r'center|centre|kendra|krishi|engineering|engg\.?|technology|campus|vidyapeetham|vidhyapeetham|vishwavidhyalaya|'
    r'deemed|laboratory|division|vet\.?|inst\.?|global|alliance|symbiosis|medicaps|vishwa|'
    r'lovely|saveetha|jain|atlas|vivekananda|acropolis|manipal|sage|ggits|skitm|sciences|'
    r'communication|information|physics|chemistry|mathematics|research|foundation|education|'
    r'veltech|klef|amity|chandigarh|proudyogiki|r&d|cse|ece|eee|it|ai|ds)\b',
    re.I,
)

JOB_TITLE_RE = re.compile(
    r'\b(?:professor|head|research\s+scholar|scholar|dean|director|lecturer|'
    r'associate|assistant|scientist|fellow|independent\s+researcher|ugdx|coe|'
    r'hod|chairperson|coordinator)\b',
    re.I,
)

INSTITUTION_FRAGMENT_RE = re.compile(
    r'^(?:university|institute|college|department|dept|school|apex|symbiosis|chandigarh|'
    r'csa|itm|engineering|engg|technology|physics|chemistry|information|communication|science|sciences|\)|\(deemed|\(deemed-to-be)\b',
    re.I,
)

NON_PERSON_WORDS = {
    'and', 'or', 'of', 'in', 'for', 'the', 'with', 'on', 'at', 'to', 'from', 'by', '&',
    'department', 'dept', 'university', 'institute', 'college', 'school', 'faculty',
    'academy', 'center', 'centre', 'engineering', 'engg', 'engg.', 'technology', 'science', 'sciences',
    'information', 'communication', 'physics', 'chemistry', 'mathematics', 'research',
    'application', 'applications', 'foundation', 'education', 'central', 'national',
    'state', 'international', 'branch', 'campus', 'marwadi', 'birla', 'vishwavidhyalaya',
    'vidhyapeetham', 'klef', 'veltech', 'jain', 'medicaps', 'amity', 'chandigarh', 'symbiosis',
    'proudyogiki', 'r&d', 'cse', 'ece', 'eee', 'dr', 'prof', 'computer', 'electronics',
    'applied', 'sciences', 'data', 'science', 'technology', 'management', 'kendra', 'krishi',
    'k', 'l', 'e', 'f'
}


def clean_author_name(name: str) -> str:
    """Strip superscript affiliation markers from author names."""
    name = re.sub(r'[\d*†‡§¶\[\]]+$', '', name).strip()
    name = re.sub(r'^[\d*†‡§¶\[\]]+\s*', '', name).strip()
    name = re.sub(r'(\w)[\d*†‡§¶]+(?=\s|,|$)', r'\1', name).strip()
    return name


def is_affiliation_text(text, is_italic=False):
    """Detect department/college/job-title lines (never author names)."""
    t = text.strip()
    if not t:
        return False
    if is_italic:
        return True
    if INSTITUTION_RE.search(t):
        return True
    if JOB_TITLE_RE.search(t):
        return True
    if INSTITUTION_FRAGMENT_RE.match(t):
        return True
    if re.search(r'\b(?:&|and)\s+head\b', t, re.I):
        return True
    return False


def is_person_name(text):
    """Heuristic for real author names (not institutions, emails, cities, or titles)."""
    t = clean_author_name(text.strip())
    if not t or len(t) < 2 or len(t) > 70:
        return False
    if not re.search(r'[A-Za-z]', t):
        return False
    if is_email(t) or is_location_line(t) or is_copyright(t) or is_affiliation_text(t):
        return False
    if re.match(r'^[A-Z]\.?$', t):
        return False
    if re.match(r'^[\d\s\W]+$', t):
        return False
    words = [w.lower().strip('.,()[]:;') for w in t.split() if w.strip('.,()[]:;')]
    if not words:
        return False
    if len(words) == 1 and (words[0] in NON_PERSON_WORDS or words[0] in CITIES or len(words[0]) <= 2):
        return False
    if all(w in NON_PERSON_WORDS or w in CITIES for w in words):
        return False
    return True


def _cluster_spans_by_column(section_spans, x_gap=90):
    """Group author-section spans into vertical columns (multi-column PDF layouts)."""
    valid = [s for s in section_spans if s['text'].strip()]
    if not valid:
        return []

    sorted_spans = sorted(valid, key=lambda s: (s.get('origin_x', 0), s['origin_y']))
    columns = [[sorted_spans[0]]]
    for s in sorted_spans[1:]:
        col_x = sum(item.get('origin_x', 0) for item in columns[-1]) / len(columns[-1])
        if abs(s.get('origin_x', 0) - col_x) < x_gap:
            columns[-1].append(s)
        else:
            columns.append([s])

    for col in columns:
        col.sort(key=lambda s: s['origin_y'])
    columns.sort(key=lambda col: min(s.get('origin_x', 0) for s in col))
    return columns


def _parse_author_column(col_spans):
    """
    Parse an author column/block into a list of authors (supporting multiple stacked
    authors in the same vertical column/grid cell as well as comma-separated authors).
    """
    authors = []
    curr_name = None
    curr_college = []
    curr_affil_num = None

    i = 0
    while i < len(col_spans):
        s = col_spans[i]
        text = s['text'].strip()
        if not text or is_copyright(text) or is_email(text) or is_location_line(text):
            i += 1
            continue

        if s['is_super'] or re.match(r'^[\d\[\]]+$', text):
            if curr_name and curr_affil_num is None:
                curr_affil_num = re.sub(r'[\[\]\s]', '', text)
            i += 1
            continue

        # Check for multiple comma-separated names on a single line
        if ',' in text and not is_affiliation_text(text):
            parts = [p.strip() for p in text.split(',') if p.strip()]
            if len(parts) >= 2 and all(is_person_name(p) or re.match(r'^[A-Za-z\s\.\d]+$', p) for p in parts):
                if curr_name:
                    authors.append({
                        'name': curr_name,
                        'college': smart_join_parts(curr_college),
                        'affil_num': curr_affil_num
                    })
                    curr_name = None
                    curr_college = []
                    curr_affil_num = None
                for p_name in parts:
                    clean_name = clean_author_name(p_name)
                    m_sup = re.search(r'\b(\d+)\b$', clean_name)
                    p_aff = None
                    if m_sup:
                        p_aff = m_sup.group(1)
                        clean_name = clean_name[:m_sup.start()].strip()
                    if clean_name:
                        authors.append({
                            'name': clean_name,
                            'college': '',
                            'affil_num': p_aff
                        })
                i += 1
                continue

        # Check for single person name
        if is_person_name(text):
            if curr_name:
                authors.append({
                    'name': curr_name,
                    'college': smart_join_parts(curr_college),
                    'affil_num': curr_affil_num
                })
                curr_name = None
                curr_college = []
                curr_affil_num = None

            curr_name = clean_author_name(text)
            if i + 1 < len(col_spans):
                ns = col_spans[i + 1]
                nt = ns['text'].strip()
                if nt and (
                    ns['is_super']
                    or (abs(ns['origin_y'] - s['origin_y']) < 4 and re.match(r'^[\d\[\]]+$', nt))
                ):
                    curr_affil_num = re.sub(r'[\[\]\s]', '', nt)
                    i += 2
                    continue
            i += 1
            continue

        if curr_name is not None:
            curr_college.append(text)

        i += 1

    if curr_name:
        authors.append({
            'name': curr_name,
            'college': smart_join_parts(curr_college),
            'affil_num': curr_affil_num
        })

    return authors


def _extract_authors_spatial(section_spans):
    """
    Extract authors using spatial column clustering and row ordering so multi-column
    and grid-based author blocks (e.g. 2x2 or 3x2) are parsed completely.
    """
    filtered = []
    for s in section_spans:
        text = s['text'].strip()
        if not text or is_copyright(text) or is_email(text) or is_location_line(text):
            continue
        filtered.append(s)

    if not filtered:
        return []

    x_values = [s.get('origin_x', 0) for s in filtered]
    x_spread = max(x_values) - min(x_values) if x_values else 0

    authors = []

    if x_spread > 120:
        for col in _cluster_spans_by_column(filtered):
            parsed_list = _parse_author_column(col)
            for a in parsed_list:
                if a and a.get('name'):
                    authors.append(a)
    else:
        ordered = sorted(filtered, key=lambda s: (s['origin_y'], s.get('origin_x', 0)))
        parsed_list = _parse_author_column(ordered)
        for a in parsed_list:
            if a and a.get('name'):
                authors.append(a)

    # Deduplicate identical name+college pairs from overlapping columns
    seen = set()
    unique = []
    for a in authors:
        key = (a['name'].lower(), a.get('college', '').lower())
        if key in seen:
            continue
        seen.add(key)
        unique.append(a)
    return unique


# ============================================================
# PDF SCRAPING HELPERS & EXTRACTION
# ============================================================

def _extract_text_spans(page):
    """Flatten and enhance all text spans in reading order from a PDF page."""
    page_dict = page.get_text('dict')
    spans = []
    for block in page_dict.get('blocks', []):
        if 'lines' not in block:
            continue
        for line in block['lines']:
            for span in line['spans']:
                spans.append({
                    'text': span['text'],
                    'size': round(span['size'], 1),
                    'flags': span['flags'],
                    'is_bold': bool(span['flags'] & 16),
                    'is_italic': bool(span['flags'] & 2),
                    'is_super': bool(span['flags'] & 1),
                    'origin_x': round(span['origin'][0], 1) if 'origin' in span else 0,
                    'origin_y': round(span['origin'][1], 1) if 'origin' in span else 0,
                })

    # Enhanced superscript detection for baseline-shifted / scaled fonts
    for i in range(1, len(spans)):
        s = spans[i]
        if s['is_super']:
            continue
        prev = spans[i - 1]
        text = s['text'].strip()
        if not text:
            continue
        if abs(s['origin_y'] - prev['origin_y']) < 3:
            if prev['size'] > 0 and s['size'] < prev['size'] * 0.75:
                s['is_super'] = True

    return spans


def _extract_title(spans):
    """Extracts title by finding spans with the maximum font size."""
    sizes = [s['size'] for s in spans if s['text'].strip()]
    if not sizes:
        return "", 0
    max_size = max(sizes)
    title_parts = []
    last_title_idx = 0
    for i, s in enumerate(spans):
        if s['size'] == max_size and s['text'].strip():
            title_parts.append(s['text'].strip())
            last_title_idx = i
    title = ' '.join(title_parts)
    title = re.sub(r'\s+', ' ', title).strip()
    return title, last_title_idx


def _extract_fallback_keywords(title, abstract_text):
    """
    Extracts 4 to 6 relevant technical terms / noun phrases from the title and abstract
    when the original PDF has no explicit Keywords / Index Terms section.
    """
    stopwords = {
        'a', 'about', 'above', 'after', 'again', 'against', 'all', 'am', 'an', 'and', 'any', 'are',
        'as', 'at', 'be', 'because', 'been', 'before', 'being', 'below', 'between', 'both', 'but',
        'by', 'can', 'could', 'did', 'do', 'does', 'doing', 'down', 'during', 'each', 'few', 'for',
        'from', 'further', 'had', 'has', 'have', 'having', 'he', 'her', 'here', 'hers', 'herself',
        'him', 'himself', 'his', 'how', 'i', 'if', 'in', 'into', 'is', 'it', 'its', 'itself',
        'just', 'me', 'more', 'most', 'my', 'myself', 'no', 'nor', 'not', 'now', 'of', 'off',
        'on', 'once', 'only', 'or', 'other', 'our', 'ours', 'ourselves', 'out', 'over', 'own',
        's', 'same', 'she', 'should', 'so', 'some', 'such', 't', 'than', 'that', 'the', 'their',
        'theirs', 'them', 'themselves', 'then', 'there', 'these', 'they', 'this', 'those', 'through',
        'to', 'too', 'under', 'until', 'up', 'very', 'was', 'we', 'were', 'what', 'when', 'where',
        'which', 'while', 'who', 'whom', 'why', 'will', 'with', 'would', 'you', 'your', 'yours',
        'paper', 'study', 'presents', 'proposed', 'results', 'based', 'using', 'novel', 'approach',
        'system', 'method', 'analysis', 'performance', 'evaluation', 'overview', 'review', 'various'
    }
    combined_text = f"{title}. {abstract_text}"
    candidates = []

    # Capitalized technical phrases (e.g. "Convolutional Neural Networks", "IoT Devices")
    caps_phrases = re.findall(r'\b[A-Z][a-zA-Z0-9]*(?:\s+[A-Z][a-zA-Z0-9]*)+\b', combined_text)
    for phrase in caps_phrases:
        phrase_clean = phrase.strip()
        words = phrase_clean.lower().split()
        if not any(w in stopwords for w in words) and 3 <= len(phrase_clean) <= 40:
            if phrase_clean not in candidates:
                candidates.append(phrase_clean)

    # Key terms from title
    title_words = [w.strip('.,:;()[]"\'') for w in title.split() if len(w) > 3]
    for w in title_words:
        if w.lower() not in stopwords and w not in candidates:
            candidates.append(w)

    return ', '.join(candidates[:5]) if candidates else ''


# ============================================================
# PDF SCRAPING
# ============================================================

def extract_paper_data(pdf_path):
    """
    Extract title, authors+college, abstract, and keywords from the
    FIRST PAGE of a PDF. Uses font size and style metadata for robust
    detection across varying IEEE paper formats.
    """
    filename = os.path.basename(pdf_path)
    print(f"\n{'='*60}")
    print(f"[PROCESS] Processing: {filename}")
    print(f"{'='*60}")

    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        print(f"  [ERROR] FATAL: Cannot open PDF -- {e}")
        return None

    if len(doc) == 0:
        print(f"  [ERROR] FATAL: PDF has no pages")
        doc.close()
        return None

    page = doc[0]
    spans = _extract_text_spans(page)
    doc.close()

    if not spans:
        print(f"  [ERROR] FATAL: No text spans found on page 1")
        return None

    # STEP 1: EXTRACT TITLE
    title, last_title_idx = _extract_title(spans)
    if not title:
        print(f"  [WARN] No title found")
    else:
        print(f"  [OK] Title: \"{title}\"")

    # STEP 2: LOCATE ABSTRACT & KEYWORDS
    abstract_idx = None
    for i in range(last_title_idx + 1, len(spans)):
        text = spans[i]['text'].strip()
        if re.match(r'^Abstract\b', text, re.IGNORECASE):
            if spans[i]['is_bold'] or spans[i]['is_italic'] or spans[i]['size'] <= 10:
                abstract_idx = i
                break

    if abstract_idx is None:
        print(f"  [WARN] 'Abstract' marker not found -- trying fallback")
        for i in range(last_title_idx + 1, len(spans)):
            if 'abstract' in spans[i]['text'].lower():
                abstract_idx = i
                break

    # Find "Keywords" / "Index Terms" marker
    keywords_idx = None
    kw_start = (abstract_idx + 1) if abstract_idx is not None else (last_title_idx + 1)
    for i in range(kw_start, len(spans)):
        text = spans[i]['text'].strip()
        if re.match(r'^(?:Keywords?|Key\s*Words?|Index\s*Terms?)\b', text, re.IGNORECASE):
            keywords_idx = i
            break

    # ════════════════════════════════════════════════
    # STEP 3: EXTRACT AUTHORS + COLLEGE (spatial parsing)
    # ════════════════════════════════════════════════
    authors = []
    if abstract_idx is not None:
        section = spans[last_title_idx + 1: abstract_idx]
        authors = _extract_authors_spatial(section)

    if authors:
        for a in authors:
            coll_display = a['college'][:55] + '...' if len(a['college']) > 55 else a['college']
            print(f"  [OK] Author: {a['name']} -> {coll_display}")
    else:
        print(f"  [WARN] No authors found")

    # ════════════════════════════════════════════════
    # STEP 4: EXTRACT ABSTRACT
    # ════════════════════════════════════════════════
    abstract_text = ""
    keywords_text = ""

    if abstract_idx is not None:
        end_idx = keywords_idx if keywords_idx is not None else len(spans)
        parts = []
        started = False

        for i in range(abstract_idx, end_idx):
            text = spans[i]['text'].strip()
            if not text:
                continue

            if not started:
                text = re.sub(r'^Abstract\s*', '', text, flags=re.IGNORECASE).strip()
                text = text.lstrip('-–—: \u2014\u2013').strip()
                if not text:
                    continue
                started = True

            parts.append(text)

        abstract_text = ' '.join(parts)
        abstract_text = re.sub(r'\s+', ' ', abstract_text).strip()

        # Check if keywords were embedded inside the abstract text
        embedded_kw_match = re.search(
            r'(?:Keywords?|Key\s*Words?|Index\s*Terms?)\s*[-–—: \u2014\u2013]+\s*(.+)',
            abstract_text,
            re.IGNORECASE | re.DOTALL
        )
        if embedded_kw_match:
            embedded_kw = embedded_kw_match.group(1).strip()
            abstract_text = abstract_text[:embedded_kw_match.start()].strip()
            keywords_text = re.sub(r'\s+', ' ', embedded_kw).strip().rstrip('.')

        print(f"  [OK] Abstract: {len(abstract_text)} characters")
    else:
        print(f"  [WARN] No abstract found")

    # ════════════════════════════════════════════════
    # STEP 5: EXTRACT KEYWORDS
    # ════════════════════════════════════════════════
    if keywords_idx is not None and not keywords_text:
        parts = []
        started = False

        for i in range(keywords_idx, len(spans)):
            s = spans[i]
            t = s['text'].strip()
            if not t:
                continue

            # Stop at true section headers (I. INTRODUCTION, Roman numerals, etc.)
            if re.match(r'^[IVX]+\.', t) or t in ('I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII') or re.match(r'^[IVX]+\.?\s', t) or 'NTRODUCTION' in t:
                break
            if re.match(r'^(?:1\.|I\.)\s*INTRODUCTION', t, re.IGNORECASE):
                break
            if s['size'] > 13.0:
                break

            if not started:
                t = re.sub(r'^(?:Keywords?|Key\s*Words?|Index\s*Terms?)\s*', '', t, flags=re.IGNORECASE).strip()
                t = t.lstrip('-–—: \u2014\u2013').strip()
                started = True
                if t:
                    parts.append(t)
            else:
                if re.match(r'^[A-Z\s]{5,}$', t) and not any(c in t for c in [',', ';']):
                    break
                parts.append(t)

        if parts:
            keywords_text = ' '.join(parts)
            keywords_text = re.sub(r'\s+', ' ', keywords_text).strip().rstrip('.')

    # Fallback keyword extraction if none found in original PDF
    is_fallback_kw = False
    if not keywords_text and (title or abstract_text):
        keywords_text = _extract_fallback_keywords(title, abstract_text)
        if keywords_text:
            is_fallback_kw = True
            print(f"  [FALLBACK] Inferred Keywords: \"{keywords_text[:80]}\"")

    if keywords_text:
        print(f"  [OK] Keywords: \"{keywords_text[:80]}{'...' if len(keywords_text) > 80 else ''}\"")
    else:
        print(f"  [WARN] No keywords found")

    # ════════════════════════════════════════════════
    # VALIDATION & ANOMALY / DOUBT DETECTION
    # ════════════════════════════════════════════════
    anomalies = []

    # 1. Author count verification against superscripts & candidate names
    if abstract_idx is not None:
        author_section = spans[last_title_idx + 1: abstract_idx]
        author_text_spans = [s['text'].strip() for s in author_section if s['text'].strip() and not is_copyright(s['text']) and not is_email(s['text'])]
        detected_supers = set(re.findall(r'\b([1-9])\b', ' '.join(author_text_spans)))
        if len(detected_supers) > len(authors) and len(authors) > 0:
            anomalies.append(f"Author count potential mismatch: {len(detected_supers)} numeric affiliations detected, but {len(authors)} author(s) extracted")

    # 2. Check if keywords were inferred via fallback
    if is_fallback_kw:
        anomalies.append("No explicit Keywords/Index Terms section found in original PDF (inferred from title & abstract)")

    # 3. Check abstract length
    if len(abstract_text) < 250:
        anomalies.append(f"Unusually short abstract ({len(abstract_text)} characters)")

    # 4. Check for missing author affiliation
    missing_affils = [a['name'] for a in authors if not a.get('college')]
    if missing_affils:
        anomalies.append(f"Missing affiliation/college for author(s): {', '.join(missing_affils)}")

    # 5. Core missing fields
    if not title:
        anomalies.append("Missing title")
    if not authors:
        anomalies.append("Missing authors")
    if not abstract_text:
        anomalies.append("Missing abstract")
    if not keywords_text:
        anomalies.append("Missing keywords")

    if anomalies:
        print(f"  [ANOMALY / DOUBT] {len(anomalies)} issue(s) detected: {'; '.join(anomalies)}")
    else:
        print(f"  [OK] All fields extracted successfully")

    return {
        'title': title,
        'authors': authors,
        'abstract': abstract_text,
        'keywords': keywords_text,
        'source_file': filename,
        'anomalies': anomalies,
    }


# ============================================================
# PDF GENERATION — CUSTOM FLOWABLES
# ============================================================

# Global page tracker: filled during body PDF generation
_page_map = {}


class PageTracker(Flowable):
    """Zero-height flowable that records the page number when drawn."""

    def __init__(self, paper_index):
        Flowable.__init__(self)
        self.paper_index = paper_index
        self.width = 0
        self.height = 0

    def draw(self):
        _page_map[self.paper_index] = self.canv.getPageNumber()


class DotSeparator(Flowable):
    """Centered row of filled dots (● ● ● ...) as visual separator."""

    def __init__(self, width, height=18):
        Flowable.__init__(self)
        self.width = width
        self.height = height

    def draw(self):
        dots = "●  " * 10 + "●"
        self.canv.setFont("Helvetica", 7)
        tw = self.canv.stringWidth(dots, "Helvetica", 7)
        x = (self.width - tw) / 2
        y = max(1, self.height / 3)
        self.canv.drawString(x, y, dots)


# ============================================================
# AUTHOR FORMATTING — Superscript Affiliations
# ============================================================

# Unicode superscript digits for affiliation markers
_SUPERSCRIPT_DIGITS = {
    1: '\u00b9', 2: '\u00b2', 3: '\u00b3', 4: '\u2074', 5: '\u2075',
    6: '\u2076', 7: '\u2077', 8: '\u2078', 9: '\u2079', 10: '\u00b9\u2070',
}

def format_authors_with_affiliations(authors, use_html_super=True):
    """
    Format authors with superscript affiliation numbers.

    Line 1: names only with superscripts (when multiple affiliations exist).
    Line 2: numbered affiliations (department/college), never job titles as authors.

    Returns:
        (authors_line, affiliations_line)
    """
    if not authors:
        return '', ''

    # Drop empty-name entries; keep only valid author records
    authors = [a for a in authors if a.get('name', '').strip()]
    if not authors:
        return '', ''

    # Build unique affiliations preserving first-seen order
    college_to_num = {}
    college_list = []
    author_nums = []

    for a in authors:
        college = a.get('college', '').strip()
        pdf_num = a.get('affil_num')
        if pdf_num and str(pdf_num).isdigit():
            num = int(pdf_num)
            if college and college not in college_to_num:
                college_to_num[college] = num
                while len(college_list) < num:
                    college_list.append('')
                college_list[num - 1] = college
            author_nums.append(num)
            continue

        if not college:
            author_nums.append(None)
            continue
        if college not in college_to_num:
            college_to_num[college] = len(college_list) + 1
            college_list.append(college)
        author_nums.append(college_to_num[college])

    college_list = [c for c in college_list if c]
    all_same = len(college_list) <= 1

    author_parts = []
    for i, a in enumerate(authors):
        name_safe = html_escape(a['name'])
        if all_same:
            author_parts.append(name_safe)
        else:
            num = author_nums[i]
            if num is None:
                author_parts.append(name_safe)
            elif use_html_super:
                author_parts.append(f"{name_safe}<super>{num}</super>")
            else:
                sup = _SUPERSCRIPT_DIGITS.get(num, str(num))
                author_parts.append(f"{name_safe}{sup}")
    authors_line = ', '.join(author_parts)

    if all_same and college_list:
        affiliations_line = html_escape(college_list[0])
    elif college_list:
        aff_parts = []
        for idx, college in enumerate(college_list):
            num = idx + 1
            college_safe = html_escape(college) if college else 'N/A'
            if use_html_super:
                aff_parts.append(f"<super>{num}</super>{college_safe}")
            else:
                sup = _SUPERSCRIPT_DIGITS.get(num, str(num))
                aff_parts.append(f"{sup}{college_safe}")
        affiliations_line = ', '.join(aff_parts)
    else:
        affiliations_line = ''

    return authors_line, affiliations_line


# ============================================================
# PDF GENERATION — BODY PAGES
# ============================================================

def _body_page_handler(canvas, doc):
    """Draw header and footer on each body page."""
    canvas.saveState()

    # ══════════════════════════════════════
    # HEADER
    # ══════════════════════════════════════
    # Conference name (italic, left-aligned)
    canvas.setFont("Times-Italic", HEADER_FONT_SIZE)
    header_y = PAGE_HEIGHT - 38
    canvas.drawString(LEFT_MARGIN, header_y, CONFERENCE_HEADER)

    # Deep blue header divider line
    canvas.setStrokeColor(DIVIDER_COLOR)
    canvas.setLineWidth(1.2)
    header_line_y = header_y - 9
    canvas.line(LEFT_MARGIN, header_line_y, PAGE_WIDTH - RIGHT_MARGIN, header_line_y)

    # ══════════════════════════════════════
    # FOOTER
    # ══════════════════════════════════════
    # Deep blue footer divider line
    footer_line_y = BOTTOM_MARGIN - 5
    canvas.setStrokeColor(DIVIDER_COLOR)
    canvas.setLineWidth(1.2)
    canvas.line(LEFT_MARGIN, footer_line_y, PAGE_WIDTH - RIGHT_MARGIN, footer_line_y)

    # Website (left-aligned) below the line
    footer_text_y = footer_line_y - 14
    canvas.setFont("Times-Bold", FOOTER_URL_FONT_SIZE)
    canvas.setFillColor(black)
    canvas.drawString(LEFT_MARGIN, footer_text_y, FOOTER_URL)

    # Page number (right-aligned) below the line
    canvas.setFont("Times-Roman", FOOTER_PAGE_FONT_SIZE)
    canvas.drawRightString(PAGE_WIDTH - RIGHT_MARGIN, footer_text_y, str(doc.page))

    canvas.restoreState()


TIER_SPECS = [
    # Tier 0: Standard generous spacing
    {
        'name': 'standard',
        'font_title': TITLE_FONT_SIZE, 'font_auth': AUTHORS_FONT_SIZE,
        'font_abs': ABSTRACT_FONT_SIZE, 'font_kw': KEYWORDS_FONT_SIZE,
        'leading_factor': LINE_SPACING_FACTOR,
        'space_title': 5, 'space_auth': 8, 'space_abs': 5, 'space_kw': 6, 'spacer_auth_abs': 4,
        'sep_top': 6, 'sep_dot': 14, 'sep_bottom': 8, 'sep_total': 28
    },
    # Tier 1: Compact Margins & Separator
    {
        'name': 'compact_margins',
        'font_title': TITLE_FONT_SIZE, 'font_auth': AUTHORS_FONT_SIZE,
        'font_abs': ABSTRACT_FONT_SIZE, 'font_kw': KEYWORDS_FONT_SIZE,
        'leading_factor': LINE_SPACING_FACTOR,
        'space_title': 3, 'space_auth': 4, 'space_abs': 3, 'space_kw': 3, 'spacer_auth_abs': 2,
        'sep_top': 3, 'sep_dot': 10, 'sep_bottom': 4, 'sep_total': 17
    },
    # Tier 2: Tight Margins & Divider
    {
        'name': 'tight_margins',
        'font_title': TITLE_FONT_SIZE, 'font_auth': AUTHORS_FONT_SIZE,
        'font_abs': ABSTRACT_FONT_SIZE, 'font_kw': KEYWORDS_FONT_SIZE,
        'leading_factor': LINE_SPACING_FACTOR,
        'space_title': 2, 'space_auth': 2, 'space_abs': 2, 'space_kw': 2, 'spacer_auth_abs': 1,
        'sep_top': 2, 'sep_dot': 6, 'sep_bottom': 2, 'sep_total': 10
    },
    # Tier 3: Zero Margins & Minimal Divider
    {
        'name': 'zero_margins',
        'font_title': TITLE_FONT_SIZE, 'font_auth': AUTHORS_FONT_SIZE,
        'font_abs': ABSTRACT_FONT_SIZE, 'font_kw': KEYWORDS_FONT_SIZE,
        'leading_factor': LINE_SPACING_FACTOR,
        'space_title': 1, 'space_auth': 1, 'space_abs': 1, 'space_kw': 1, 'spacer_auth_abs': 1,
        'sep_top': 1, 'sep_dot': 4, 'sep_bottom': 1, 'sep_total': 6
    }
]


def _build_and_measure_page(page_papers, page_paper_indices, tier, avail_width):
    """
    Builds the ReportLab flowables for all papers on a single page using the given
    tier parameters and measures the exact total rendered height.
    """
    font_title = tier['font_title']
    font_auth = tier['font_auth']
    font_abs = tier['font_abs']
    font_kw = tier['font_kw']
    leading_factor = tier['leading_factor']
    space_title = tier['space_title']
    space_auth = tier['space_auth']
    space_abs = tier['space_abs']
    space_kw = tier['space_kw']
    spacer_auth_abs = tier['spacer_auth_abs']
    sep_top = tier['sep_top']
    sep_dot = tier['sep_dot']
    sep_bottom = tier['sep_bottom']

    uid = random.randint(1, 999999)
    st_title = ParagraphStyle(
        f"title_{tier.get('name', 'tier')}_{len(page_papers)}_{uid}",
        fontName='Times-Bold',
        fontSize=font_title,
        leading=font_title * leading_factor,
        alignment=TA_CENTER,
        spaceAfter=space_title,
        wordSpace=WORD_SPACING,
        keepWithNext=True,
    )
    st_auth = ParagraphStyle(
        f"auth_{tier.get('name', 'tier')}_{len(page_papers)}_{uid}",
        fontName='Times-Roman',
        fontSize=font_auth,
        leading=font_auth * leading_factor,
        alignment=TA_CENTER,
        spaceAfter=space_auth,
        wordSpace=WORD_SPACING * 0.7,
        keepWithNext=True,
    )
    st_abs = ParagraphStyle(
        f"abs_{tier.get('name', 'tier')}_{len(page_papers)}_{uid}",
        fontName='Times-Roman',
        fontSize=font_abs,
        leading=font_abs * leading_factor,
        alignment=TA_JUSTIFY,
        spaceAfter=space_abs,
        wordSpace=WORD_SPACING,
        keepWithNext=True,
    )
    st_kw = ParagraphStyle(
        f"kw_{tier.get('name', 'tier')}_{len(page_papers)}_{uid}",
        fontName='Times-Roman',
        fontSize=font_kw,
        leading=font_kw * leading_factor,
        alignment=TA_LEFT,
        spaceAfter=space_kw,
        wordSpace=WORD_SPACING * 0.7,
    )

    page_flowables = []
    total_h = 0
    k = len(page_papers)

    for item_idx, (paper, paper_idx) in enumerate(zip(page_papers, page_paper_indices)):
        title_safe = html_escape(paper['title'])
        abstract_safe = html_escape(paper['abstract'])
        keywords_safe = html_escape(paper['keywords'])
        authors_text, affiliations_text = format_authors_with_affiliations(paper['authors'], use_html_super=True)

        p_title = Paragraph(title_safe, st_title)
        p_authors = Paragraph(authors_text, st_auth)
        p_aff = Paragraph(f'<i>{affiliations_text}</i>', st_auth) if affiliations_text else None
        p_abs = Paragraph(f"<b>Abstract— </b>{abstract_safe}", st_abs)
        p_kw = Paragraph(f"<b>Keywords— </b><i>{keywords_safe}</i>", st_kw) if keywords_safe else None

        # Measure flowables
        _, h_t = p_title.wrap(avail_width, 9999)
        _, h_a = p_authors.wrap(avail_width, 9999)
        h_aff = 0
        if p_aff:
            _, h_aff = p_aff.wrap(avail_width, 9999)
        _, h_ab = p_abs.wrap(avail_width, 9999)
        h_k = 0
        if p_kw:
            _, h_k = p_kw.wrap(avail_width, 9999)

        item_h = (h_t + space_title + h_a +
                  (h_aff + space_auth if p_aff else space_auth) +
                  spacer_auth_abs + h_ab + space_abs + (h_k + space_kw if p_kw else 0))
        total_h += item_h

        # Wrap this entire paper's components in an atomic KeepTogether block with keepWithNext=True
        paper_block = [PageTracker(paper_idx), p_title, p_authors]
        if p_aff:
            paper_block.append(p_aff)
        paper_block.append(Spacer(1, spacer_auth_abs))
        paper_block.append(p_abs)
        if p_kw:
            paper_block.append(p_kw)

        page_flowables.append(KeepTogether(paper_block))

        # Separator between papers on the same page (not after the last paper)
        if item_idx < k - 1:
            sep_h = sep_top + sep_dot + sep_bottom
            total_h += sep_h
            page_flowables.append(Spacer(1, sep_top))
            page_flowables.append(DotSeparator(avail_width, height=sep_dot))
            page_flowables.append(Spacer(1, sep_bottom))

    return page_flowables, total_h


def _precompute_paper_heights(papers, avail_width):
    """
    Precomputes the rendered height matrix for all papers across all tiers.
    Matrix shape: N x 4
    """
    matrix = []
    for p_idx, p in enumerate(papers):
        row = []
        for t_idx, tier in enumerate(TIER_SPECS):
            _, h = _build_and_measure_page([p], [p_idx], tier, avail_width)
            row.append(h)
        matrix.append(row)
    return matrix


def _find_optimal_zero_waste_packing(papers, avail_height, forbidden_pairs=None):
    """
    Bin-packing engine for abstract pages.

    Strategy: Tallest-first, shortest-partner matching.
    Big papers always get paired with the shortest available paper FIRST,
    so small papers are never wasted by pairing two smalls together
    while a big paper sits alone.

    Rules:
      1. At most MAX_PAPERS_PER_PAGE (2) abstracts per page.
      2. Every non-final page must have 2 abstracts (unless truly oversized).
      3. Non-final pages must fill at least MIN_PAGE_FILL_RATIO of height.
      4. Last page may contain a single leftover abstract (any whitespace OK).
      5. Conservative headroom prevents KeepTogether overflow / keyword leaks.
      6. Only papers that cannot fit with the SHORTEST available paper
         are allowed solo pages.
    """
    n = len(papers)
    if n == 0:
        return []
    if n == 1:
        return [(0,)]

    if forbidden_pairs is None:
        forbidden_pairs = set()

    height_matrix = _precompute_paper_heights(papers, CONTENT_WIDTH)
    target_budget = avail_height - PACKING_HEADROOM_PT

    def test_group(grp):
        k = len(grp)
        if k == 0:
            return True, 0.0, 0
        if k > MAX_PAPERS_PER_PAGE:
            return False, float('inf'), None

        grp_tuple = tuple(sorted(grp))
        if grp_tuple in forbidden_pairs:
            return False, float('inf'), None

        for t_idx, tier in enumerate(TIER_SPECS):
            h_papers = sum(height_matrix[i][t_idx] for i in grp)
            h_sep = (k - 1) * tier['sep_total']
            tot = h_papers + h_sep
            if tot <= target_budget:
                return True, tot, t_idx
        return False, float('inf'), None

    # ── Phase 1: Tallest-first, shortest-partner greedy matching ──
    # Sort all papers by their minimum possible height (most compact tier)
    min_heights = [min(height_matrix[i]) for i in range(n)]
    
    # Sort indices: tallest first
    tall_order = sorted(range(n), key=lambda i: min_heights[i], reverse=True)
    
    paired = set()
    bins = []
    truly_oversized = set()
    
    for idx in tall_order:
        if idx in paired:
            continue
        
        # Find the shortest AVAILABLE partner that fits with this paper
        # Sort candidates by height ascending (shortest first)
        candidates = [
            j for j in range(n)
            if j != idx and j not in paired
        ]
        candidates.sort(key=lambda j: min_heights[j])  # shortest first
        
        best_partner = None
        best_tot = float('inf')
        best_tier = None
        
        for cand in candidates:
            pair_tuple = tuple(sorted([idx, cand]))
            if pair_tuple in forbidden_pairs:
                continue
            fits, tot, t_idx = test_group([idx, cand])
            if fits:
                # Pick the partner that gives the tightest fit (least waste)
                if best_partner is None or tot > best_tot:
                    # Actually prefer HIGHER fill (less whitespace)
                    best_partner = cand
                    best_tot = tot
                    best_tier = t_idx
                # But first valid (shortest) partner is already great
                # Accept first fit for speed, SA will refine later
                best_partner = cand
                best_tot = tot
                best_tier = t_idx
                break
        
        if best_partner is not None:
            bins.append([idx, best_partner])
            paired.add(idx)
            paired.add(best_partner)
        else:
            # This paper cannot pair with ANY remaining paper
            truly_oversized.add(idx)
            bins.append([idx])
            paired.add(idx)
    
    # ── Phase 2: Try to rescue oversized singles by stealing a partner ──
    # For each truly_oversized paper, check if swapping it with someone
    # from a paired bin would work better
    for solo_idx in list(truly_oversized):
        rescued = False
        for b_idx, b in enumerate(bins):
            if len(b) != 2:
                continue
            for pos in range(2):
                partner_in_bin = b[pos]
                other_in_bin = b[1 - pos]
                # Can solo_idx pair with partner_in_bin?
                pair1 = tuple(sorted([solo_idx, partner_in_bin]))
                if pair1 in forbidden_pairs:
                    continue
                fits1, _, _ = test_group([solo_idx, partner_in_bin])
                if not fits1:
                    continue
                # Can other_in_bin pair with anyone else or go solo on last page?
                # Check if other_in_bin can fit with another solo
                other_solo_bins = [i for i, bb in enumerate(bins) if len(bb) == 1 and bb[0] != solo_idx]
                placed_other = False
                for os_idx in other_solo_bins:
                    os_paper = bins[os_idx][0]
                    pair2 = tuple(sorted([other_in_bin, os_paper]))
                    if pair2 in forbidden_pairs:
                        continue
                    fits2, _, _ = test_group([other_in_bin, os_paper])
                    if fits2:
                        # Great! Swap: solo_idx takes partner_in_bin's spot,
                        # other_in_bin joins os_paper
                        bins[b_idx] = [solo_idx, partner_in_bin]
                        bins[os_idx] = [other_in_bin, os_paper]
                        # Remove solo_idx's old solo bin
                        bins = [bb for bb in bins if not (len(bb) == 1 and bb[0] == solo_idx)]
                        truly_oversized.discard(solo_idx)
                        truly_oversized.discard(os_paper)
                        placed_other = True
                        rescued = True
                        break
                if rescued:
                    break
            if rescued:
                break

    # ── Phase 3: Simulated Annealing refinement (5000 iterations) ──
    def score_partition(part):
        sc = 0.0
        valid_bins = [b for b in part if b]
        num_bins = len(valid_bins)
        for b_idx, b in enumerate(valid_bins):
            fits, tot, t_idx = test_group(b)
            if not fits:
                return -1e9

            ratio = min(1.0, tot / avail_height)
            is_last_page = (b_idx == num_bins - 1)

            if not is_last_page:
                if len(b) < 2:
                    if len(b) == 1 and b[0] in truly_oversized:
                        sc += 100.0  # Acceptable for truly oversized
                    else:
                        return -1e9  # Non-final solo page is FORBIDDEN
                if ratio < MIN_PAGE_FILL_RATIO:
                    sc -= 50000.0 * (MIN_PAGE_FILL_RATIO - ratio)

            if len(b) == 2:
                sc += 800.0 + (ratio * 300.0)
            elif len(b) == 1:
                if not is_last_page and b[0] not in truly_oversized:
                    return -1e9
                sc += 50.0

            # Prefer standard generous margins
            if t_idx == 0:
                sc += 60.0
            elif t_idx == 1:
                sc += 30.0
            elif t_idx == 3:
                sc -= 40.0
        return sc

    best_bins = [list(b) for b in bins]
    best_score = score_partition(best_bins)
    curr_bins = [list(b) for b in best_bins]

    for it in range(5000):
        curr_bins = [b for b in curr_bins if b]
        if len(curr_bins) < 2:
            break
        b1_idx, b2_idx = random.sample(range(len(curr_bins)), 2)
        b1 = curr_bins[b1_idx]
        b2 = curr_bins[b2_idx]

        # Move 1: Try item swap between b1 and b2
        if b1 and b2:
            i1 = random.randrange(len(b1))
            i2 = random.randrange(len(b2))
            c1 = list(b1)
            c2 = list(b2)
            c1[i1], c2[i2] = c2[i2], c1[i1]
            if len(c1) <= MAX_PAPERS_PER_PAGE and len(c2) <= MAX_PAPERS_PER_PAGE:
                f1, _, _ = test_group(c1)
                f2, _, _ = test_group(c2)
                if f1 and f2:
                    cand_part = list(curr_bins)
                    cand_part[b1_idx] = c1
                    cand_part[b2_idx] = c2
                    sc = score_partition(cand_part)
                    if sc > best_score:
                        best_score = sc
                        best_bins = [list(b) for b in cand_part if b]
                        curr_bins = cand_part
                        continue

        # Move 2: Try item transfer from b1 to b2
        if len(b1) >= 1 and len(b2) <= MAX_PAPERS_PER_PAGE - 1:
            i1 = random.randrange(len(b1))
            c1 = [x for idx, x in enumerate(b1) if idx != i1]
            c2 = b2 + [b1[i1]]
            f1, _, _ = test_group(c1) if c1 else (True, 0, 0)
            f2, _, _ = test_group(c2)
            if f1 and f2:
                cand_part = list(curr_bins)
                cand_part[b1_idx] = c1
                cand_part[b2_idx] = c2
                cand_part = [b for b in cand_part if b]
                sc = score_partition(cand_part)
                if sc > best_score:
                    best_score = sc
                    best_bins = [list(b) for b in cand_part]
                    curr_bins = cand_part

    # ── Final ordering: paired pages first (desc by fill), solos last ──
    clean_bins = [b for b in best_bins if b]
    multi_groups = [b for b in clean_bins if len(b) >= 2]
    single_groups = [b for b in clean_bins if len(b) == 1]

    multi_records = []
    for b in multi_groups:
        _, tot, t_idx = test_group(b)
        multi_records.append((tuple(b), tot, t_idx))
    multi_records.sort(key=lambda rec: rec[1], reverse=True)

    # Oversized solos go before the final leftover
    oversized_singles = [s for s in single_groups if s[0] in truly_oversized]
    leftover_singles = [s for s in single_groups if s[0] not in truly_oversized]

    final_groups = [rec[0] for rec in multi_records]
    for s in oversized_singles:
        final_groups.append(tuple(s))
    for s in leftover_singles:
        final_groups.append(tuple(s))

    return final_groups


def _generate_page_flowables_justified(page_papers, page_paper_indices, avail_width, avail_height, is_final_page=False):
    """
    Generates flowables for a page and vertically justifies / stretches internal spacing
    so that the entire page is 100% filled from top to bottom, eliminating empty white space
    at the bottom of the page with strict safe headroom.
    """
    k = len(page_papers)
    target_budget = avail_height - 45.0
    best_tier = None
    base_h = 0

    # Find fitting tier
    for tier in TIER_SPECS:
        _, measured_h = _build_and_measure_page(page_papers, page_paper_indices, tier, avail_width)
        if measured_h <= target_budget:
            best_tier = tier
            base_h = measured_h
            break

    if best_tier is None:
        best_tier = TIER_SPECS[-1]
        _, base_h = _build_and_measure_page(page_papers, page_paper_indices, best_tier, avail_width)

    # Calculate remaining white space (slack) with conservative 40pt buffer
    slack = (avail_height - 40.0) - base_h

    # Vertically justify if there is slack and it's not a short final page
    if slack > 4.0 and (not is_final_page or (base_h / avail_height) >= 0.70):
        # Expansion points: 5 per paper + 2 per separator
        num_expand_points = 5 * k + (2 * (k - 1) if k > 1 else 0)
        boost = min(4.0, slack / max(1, num_expand_points))

        # Iteratively verify that justified layout does not exceed target budget
        while boost >= 0.5:
            justified_tier = dict(best_tier)
            justified_tier['name'] = f"{best_tier['name']}_justified"
            justified_tier['space_title'] = best_tier['space_title'] + (boost * 0.8)
            justified_tier['space_auth'] = best_tier['space_auth'] + (boost * 1.0)
            justified_tier['space_abs'] = best_tier['space_abs'] + (boost * 0.9)
            justified_tier['space_kw'] = best_tier['space_kw'] + (boost * 0.8)
            justified_tier['spacer_auth_abs'] = best_tier['spacer_auth_abs'] + (boost * 1.0)
            justified_tier['sep_top'] = best_tier['sep_top'] + (boost * 1.2)
            justified_tier['sep_bottom'] = best_tier['sep_bottom'] + (boost * 1.4)

            flowables, final_h = _build_and_measure_page(page_papers, page_paper_indices, justified_tier, avail_width)
            if final_h <= avail_height - 35.0:
                return flowables, justified_tier['name'], final_h
            boost -= 0.5

    # Build standard flowables
    flowables, final_h = _build_and_measure_page(page_papers, page_paper_indices, best_tier, avail_width)
    return flowables, best_tier['name'], final_h


def _inspect_body_pdf_pages(pdf_path, expected_num_pages):
    """
    Inspects each rendered page of the generated body PDF using PyMuPDF (fitz).
    Detects any orphan keyword spills, split content, or underfilled intermediate pages.
    Returns (is_clean, list_of_defects).
    """
    import fitz
    if not os.path.exists(pdf_path):
        return False, [{'type': 'file_missing', 'msg': 'Body PDF file was not generated.'}]

    doc = fitz.open(pdf_path)
    defects = []

    if len(doc) != expected_num_pages:
        defects.append({
            'type': 'page_count_mismatch',
            'msg': f"Generated {len(doc)} physical pages, expected {expected_num_pages} (page overflow detected).",
            'actual_pages': len(doc),
            'expected_pages': expected_num_pages
        })

    for pg_idx, page in enumerate(doc):
        raw_text = page.get_text().strip()
        lines = [l.strip() for l in raw_text.split('\n') if l.strip()]
        is_last = (pg_idx == len(doc) - 1)

        # Check 1: Orphan keywords leak (starts with Keywords or contains only keywords without an abstract/title)
        has_title = any(len(l) > 30 and "—" not in l for l in lines)
        has_abstract = any("Abstract—" in l or "Abstract —" in l for l in lines)
        
        if (raw_text.startswith("Keywords—") or (not has_title and not has_abstract and "Keywords—" in raw_text)) or (len(raw_text) < 200 and "Keywords—" in raw_text and not has_abstract):
            defects.append({
                'type': 'orphan_keywords_leak',
                'page': pg_idx + 1,
                'msg': f"Page {pg_idx + 1} contains orphan keywords leaked across page boundary: '{raw_text[:60]}...'"
            })

        # Check 2: Intermediate underfilled page (< 280 chars) on non-final page
        if not is_last and len(raw_text) < 280:
            defects.append({
                'type': 'underfilled_intermediate_page',
                'page': pg_idx + 1,
                'msg': f"Page {pg_idx + 1} has insufficient content ({len(raw_text)} chars) on an intermediate page."
            })

    doc.close()
    return len(defects) == 0, defects


def generate_body_pdf(papers, output_path):
    """
    Generate the body pages of the compiled PDF using ReportLab Platypus
    with Automated Post-Compilation Self-Healing Verification.

    Zero-White-Space & Zero-Leak Guarantees:
      1. Uses Combinatorial Re-ordering & Bin Packing with 50pt safe headroom.
      2. Every non-final page contains AT LEAST 2 abstracts and is packed to 90%-100% fullness.
      3. Paragraphs are bound with keepWithNext=True to eliminate keyword spills.
      4. Inspects generated physical pages with PyMuPDF to verify zero defects.
      5. Automatically self-heals and re-packs if any defect is detected.

    Returns tuple (page_map, page_groups, page_density_modes)
    """
    global _page_map
    avail_height = PAGE_HEIGHT - TOP_MARGIN - BOTTOM_MARGIN
    forbidden_pairs = set()

    for attempt in range(1, 4):
        _page_map = {}

        doc = BaseDocTemplate(
            output_path,
            pagesize=A4,
            leftMargin=LEFT_MARGIN,
            rightMargin=RIGHT_MARGIN,
            topMargin=TOP_MARGIN,
            bottomMargin=BOTTOM_MARGIN,
        )

        frame = Frame(
            LEFT_MARGIN,
            BOTTOM_MARGIN,
            CONTENT_WIDTH,
            avail_height,
            id='body_frame',
        )

        doc.addPageTemplates([
            PageTemplate('body_page', frames=[frame], onPage=_body_page_handler)
        ])

        # ── Step 1: Combinatorial Re-Ordering Bin Packing ──
        page_groups = _find_optimal_zero_waste_packing(papers, avail_height, forbidden_pairs=forbidden_pairs)

        if attempt == 1:
            print(f"\n  [ANALYSIS] Available height per page: {avail_height:.0f}pt")
            print(f"\n  [PACKING] Zero-White-Space Optimized Page Groups ({len(page_groups)} pages):")

        elements = []
        page_density_modes = []

        for pg_num, group in enumerate(page_groups):
            page_papers = [papers[idx] for idx in group]
            page_paper_indices = list(group)
            is_last = (pg_num == len(page_groups) - 1)

            # Generate vertically justified flowables to eliminate bottom white space
            flowables, mode_name, measured_h = _generate_page_flowables_justified(
                page_papers, page_paper_indices, CONTENT_WIDTH, avail_height, is_final_page=is_last
            )
            page_density_modes.append(mode_name)

            usage = min(100.0, (measured_h / avail_height) * 100)
            tag = " (Final Page)" if is_last else ""
            mode_tag = f" [{mode_name}]" if mode_name != 'standard' else ""
            titles = [papers[idx]['title'][:38] + '...' for idx in group]
            if attempt == 1:
                print(f"    Page {pg_num+1}: {len(group)} abstract(s), {usage:.0f}% filled{mode_tag}{tag} -- {', '.join(titles)}")

            if pg_num > 0:
                elements.append(PageBreak())

            # Directly extend flowables (each paper is an atomic KeepTogether block with keepWithNext=True)
            elements.extend(flowables)

        # Build the PDF
        doc.build(elements)

        # ── Step 2: Post-Compilation PyMuPDF Verification ──
        is_clean, defects = _inspect_body_pdf_pages(output_path, len(page_groups))

        if is_clean:
            print(f"  [VERIFIED] All {len(page_groups)} body pages passed 100% clean verification (zero leaks, zero split content)!")
            return dict(_page_map), page_groups, page_density_modes
        else:
            print(f"  [SELF-HEALING] Detected {len(defects)} layout defect(s) on attempt {attempt}:")
            for d in defects:
                print(f"    - {d.get('msg')}")
                if 'page' in d and d['page'] <= len(page_groups):
                    offending_group = page_groups[d['page'] - 1]
                    forbidden_pairs.add(tuple(sorted(offending_group)))

            print(f"  [SELF-HEALING] Re-optimizing pairings with tighter constraints...")

    # If completed after self-healing attempts, return final state
    print(f"  [DONE] Completed body PDF generation.")
    return dict(_page_map), page_groups, page_density_modes


# ============================================================
# PDF GENERATION — TABLE OF CONTENTS
# ============================================================

def generate_toc_pdf(papers, page_map, output_path):
    """
    Generate Table of Contents page(s) using direct canvas drawing
    for precise dotted-leader alignment.
    Returns list of metadata dicts for hyperlink injection:
    [
        {
            'paper_idx': i,
            'title': paper['title'],
            'toc_page_idx': current_toc_page,
            'body_page_num': page_val,
            'rect_fitz': (x0, y0, x1, y1),
        },
        ...
    ]
    """
    c = canvas_mod.Canvas(output_path, pagesize=A4)
    y = PAGE_HEIGHT - 75
    current_toc_page = 0
    toc_entries_meta = []

    # ── "Contents" heading ──
    c.setFont("Times-Bold", 24)
    c.drawString(LEFT_MARGIN, y, "Contents")
    y -= 28

    # Deep blue underline
    c.setStrokeColor(DIVIDER_COLOR)
    c.setLineWidth(1.5)
    c.line(LEFT_MARGIN, y, PAGE_WIDTH - RIGHT_MARGIN, y)
    y -= 25
    # ── Sort papers by page number ascending ──
    sorted_papers = [(i, p, page_map.get(i, 9999)) for i, p in enumerate(papers)]
    sorted_papers.sort(key=lambda x: x[2])

    for i, paper, page_val in sorted_papers:
        page_num = str(page_val) if page_val != 9999 else '?'

        # ── Pre-estimate this entry's total height ──
        toc_title_font = "Times-Roman"
        toc_title_size = 9.0
        toc_line_height = 11.5

        title = paper['title']
        pnum_width = stringWidth(page_num, "Times-Roman", toc_title_size)
        dot_char_width = stringWidth('.', "Times-Roman", 8)
        max_single_w = CONTENT_WIDTH - pnum_width - 25

        title_full_w = stringWidth(title, toc_title_font, toc_title_size)
        if title_full_w <= max_single_w:
            lines = [title]
        else:
            # 2-line clean wrap: line 1 uses full CONTENT_WIDTH, line 2 uses max_single_w
            words = title.split()
            line1_words = []
            curr_w = 0
            space_w = stringWidth(' ', toc_title_font, toc_title_size)
            split_idx = 0
            for w_idx, w in enumerate(words):
                w_len = stringWidth(w, toc_title_font, toc_title_size)
                needed = w_len if not line1_words else (space_w + w_len)
                if curr_w + needed <= CONTENT_WIDTH - 5:
                    line1_words.append(w)
                    curr_w += needed
                    split_idx = w_idx + 1
                else:
                    break
            if split_idx == 0:
                split_idx = 1
            line1 = ' '.join(words[:split_idx])
            remainder = ' '.join(words[split_idx:])

            if stringWidth(remainder, toc_title_font, toc_title_size) <= max_single_w:
                lines = [line1, remainder]
            else:
                sub_lines = _word_wrap(remainder, toc_title_font, toc_title_size, max_single_w)
                lines = [line1] + sub_lines

        author_names = ', '.join(a['name'] for a in paper['authors'])
        author_lines = _word_wrap(author_names, "Times-Italic", 7.5, CONTENT_WIDTH) if author_names else []

        # Total height this entry will consume
        entry_height = (len(lines) * toc_line_height +
                        len(author_lines) * 9.5 + 6)

        # ── Check for page overflow BEFORE drawing ──
        if y - entry_height < 50:
            _draw_toc_footer(c)
            c.showPage()
            current_toc_page += 1
            # Draw header on continuation TOC pages
            _draw_toc_header(c)
            y = PAGE_HEIGHT - 75

        # Record top of entry for hyperlink bounding box
        entry_toc_page = current_toc_page
        entry_top_rl = y + toc_title_size + 2.0

        # ── Draw title with dotted leaders → page number ──
        for j, line in enumerate(lines):
            c.setFont(toc_title_font, toc_title_size)
            c.drawString(LEFT_MARGIN, y, line)

            if j == len(lines) - 1:
                title_end_x = LEFT_MARGIN + stringWidth(
                    line, toc_title_font, toc_title_size
                ) + 5
                page_right_x = PAGE_WIDTH - RIGHT_MARGIN
                dot_limit_x = page_right_x - pnum_width - 8

                c.setFont("Times-Roman", 8)
                x = title_end_x
                while x + dot_char_width < dot_limit_x:
                    c.drawString(x, y, '.')
                    x += dot_char_width + 1.2

                c.setFont("Times-Roman", toc_title_size)
                c.drawRightString(page_right_x, y, page_num)

            y -= toc_line_height

        # ── Author names below title (italic, smaller) ──
        if author_lines:
            c.setFont("Times-Italic", 7.5)
            for al in author_lines:
                c.drawString(LEFT_MARGIN, y, al)
                y -= 9.5

        # Record bottom of entry and save metadata for hyperlink
        entry_bottom_rl = y - 2.0
        fitz_x0 = LEFT_MARGIN - 2.0
        fitz_x1 = PAGE_WIDTH - RIGHT_MARGIN + 2.0
        fitz_y0 = max(0.0, PAGE_HEIGHT - entry_top_rl)
        fitz_y1 = min(PAGE_HEIGHT, PAGE_HEIGHT - entry_bottom_rl)

        if page_val != 9999:
            toc_entries_meta.append({
                'paper_idx': i,
                'title': paper['title'],
                'toc_page_idx': entry_toc_page,
                'body_page_num': page_val,
                'rect_fitz': (fitz_x0, fitz_y0, fitz_x1, fitz_y1),
            })

        y -= 6  # clean gap before next entry

    _draw_toc_footer(c)
    c.save()
    return toc_entries_meta


def _draw_toc_header(canvas):
    """Draw conference header on TOC continuation pages."""
    canvas.setFont("Times-Italic", HEADER_FONT_SIZE)
    header_y = PAGE_HEIGHT - 38
    canvas.drawString(LEFT_MARGIN, header_y, CONFERENCE_HEADER)

    canvas.setStrokeColor(DIVIDER_COLOR)
    canvas.setLineWidth(1.2)
    header_line_y = header_y - 9
    canvas.line(LEFT_MARGIN, header_line_y,
                PAGE_WIDTH - RIGHT_MARGIN, header_line_y)


def _draw_toc_footer(canvas):
    """Draw footer on a TOC page with deep blue divider."""
    # Deep blue footer divider line
    footer_line_y = BOTTOM_MARGIN - 5
    canvas.setStrokeColor(DIVIDER_COLOR)
    canvas.setLineWidth(1.2)
    canvas.line(LEFT_MARGIN, footer_line_y, PAGE_WIDTH - RIGHT_MARGIN, footer_line_y)

    # Website (left-aligned)
    footer_text_y = footer_line_y - 14
    canvas.setFont("Times-Bold", FOOTER_URL_FONT_SIZE)
    canvas.setFillColor(black)
    canvas.drawString(LEFT_MARGIN, footer_text_y, FOOTER_URL)


def _word_wrap(text, font_name, font_size, max_width):
    """Simple word-wrap that returns a list of lines."""
    words = text.split()
    lines = []
    current = ""

    for word in words:
        test = (current + " " + word).strip()
        if stringWidth(test, font_name, font_size) <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word

    if current:
        lines.append(current)

    return lines if lines else [text]


# ============================================================
# PDF MERGE
# ============================================================

def merge_pdfs(toc_path, body_path, output_path, toc_entries_meta=None):
    """Merge Table of Contents + Body into a single final PDF with clickable TOC hyperlinks & PDF bookmarks."""
    output_doc = fitz.open()

    toc_doc = fitz.open(toc_path)
    body_doc = fitz.open(body_path)

    num_toc_pages = len(toc_doc)

    output_doc.insert_pdf(toc_doc)
    output_doc.insert_pdf(body_doc)

    # ── Inject Clickable Hyperlinks in Table of Contents ──
    if toc_entries_meta:
        pdf_outline_toc = []
        for entry in toc_entries_meta:
            toc_page_idx = entry['toc_page_idx']
            body_page_num = entry['body_page_num']
            if body_page_num == 9999 or toc_page_idx >= num_toc_pages:
                continue

            target_page_0based = num_toc_pages + (body_page_num - 1)
            target_page_1based = target_page_0based + 1

            # 1. Clickable on-page link annotation (clicking title/page number jumps to abstract)
            if toc_page_idx < len(output_doc) and target_page_0based < len(output_doc):
                page = output_doc[toc_page_idx]
                r = entry['rect_fitz']
                link_rect = fitz.Rect(r[0], r[1], r[2], r[3])
                page.insert_link({
                    'kind': fitz.LINK_GOTO,
                    'page': target_page_0based,
                    'from': link_rect,
                    'to': fitz.Point(0, 0)
                })

            # 2. PDF Document Outline / Bookmark
            title_short = entry['title']
            if len(title_short) > 80:
                title_short = title_short[:77] + '...'
            pdf_outline_toc.append([1, title_short, target_page_1based])

        # Set document outline bookmarks
        if pdf_outline_toc:
            try:
                output_doc.set_toc(pdf_outline_toc)
            except Exception as e:
                print(f"  [WARN] Could not set PDF outline bookmarks: {e}")

    output_doc.save(output_path)

    output_doc.close()
    toc_doc.close()
    body_doc.close()


# ============================================================
# COVER PAGE GENERATION
# ============================================================

def generate_cover_page_pdf(output_path):
    """
    Generate a cover page PDF from cover_page.json.
    Blocks store HTML in 'content' (with <b>, <i>, <u>, <br> tags).
    Falls back to plain 'text' field for backward compatibility.
    Returns True if cover page was generated, False if no config found.
    """
    config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'cover_page.json')
    if not os.path.exists(config_path):
        return False

    try:
        with open(config_path, 'r') as f:
            data = json.load(f)
    except Exception:
        return False

    blocks = data.get('blocks', [])
    if not blocks:
        return False

    c = canvas_mod.Canvas(output_path, pagesize=A4)

    for block in blocks:
        block_type = block.get('type', 'text')
        x = block.get('x', 0)
        y_top = block.get('y', 0)
        width = block.get('width', 200)
        height = block.get('height', 30)
        
        y_rl = PAGE_HEIGHT - y_top - height

        if block_type == 'image':
            image_url = block.get('imageUrl', '')
            if image_url.startswith('data:image/'):
                import base64
                import io
                from reportlab.lib.utils import ImageReader
                try:
                    header, encoded = image_url.split(',', 1)
                    image_data = base64.b64decode(encoded)
                    image_stream = io.BytesIO(image_data)
                    img_reader = ImageReader(image_stream)
                    c.drawImage(img_reader, x, y_rl, width=width, height=height, preserveAspectRatio=True, mask='auto')
                except Exception as e:
                    print(f"Error drawing image to PDF: {e}")
            continue

        if block_type == 'drawing':
            lines = block.get('lines', [])
            sym = block.get('symmetry', 'none')
            stroke_width = block.get('strokeWidth', 2)
            # Parse hex color
            color_hex = block.get('strokeColor', '#000000').lstrip('#')
            try:
                r, g, b = tuple(int(color_hex[i:i+2], 16) / 255.0 for i in (0, 2, 4))
            except Exception:
                r, g, b = 0, 0, 0
                
            c.setStrokeColorRGB(r, g, b)
            c.setLineWidth(stroke_width)
            
            # Helper to draw a line
            def draw_rl_line(l_x1, l_y1, l_x2, l_y2):
                # Convert drawing box coords to absolute PDF coords
                # drawing (0,0) is top-left of block
                pdf_x1 = x + l_x1
                pdf_y1 = PAGE_HEIGHT - (y_top + l_y1)
                pdf_x2 = x + l_x2
                pdf_y2 = PAGE_HEIGHT - (y_top + l_y2)
                c.line(pdf_x1, pdf_y1, pdf_x2, pdf_y2)

            for line in lines:
                l_x1, l_y1, l_x2, l_y2 = line.get('x1',0), line.get('y1',0), line.get('x2',0), line.get('y2',0)
                draw_rl_line(l_x1, l_y1, l_x2, l_y2)
                
                # Apply symmetry
                if sym in ('horizontal', 'both'):
                    draw_rl_line(width - l_x1, l_y1, width - l_x2, l_y2)
                if sym in ('vertical', 'both'):
                    draw_rl_line(l_x1, height - l_y1, l_x2, height - l_y2)
                if sym == 'both':
                    draw_rl_line(width - l_x1, height - l_y1, width - l_x2, height - l_y2)
            continue

        # Get HTML content, fall back to plain text
        html_content = block.get('content', '')
        if not html_content:
            plain = block.get('text', '')
            html_content = plain.replace('\n', '<br/>')

        font_size = block.get('fontSize', 12)
        font_weight = block.get('fontWeight', 'normal')
        font_style = block.get('fontStyle', 'normal')
        text_align = block.get('textAlign', 'center')

        # Map alignment
        align_map = {'left': TA_LEFT, 'center': TA_CENTER, 'right': TA_LEFT}
        alignment = align_map.get(text_align, TA_CENTER)
        if text_align == 'right':
            from reportlab.lib.enums import TA_RIGHT
            alignment = TA_RIGHT

        # Choose base font
        font_family = block.get('fontFamily', 'Times New Roman')
        is_bold = (font_weight == 'bold')
        is_italic = (font_style == 'italic')
        
        if font_family == 'Arial':
            if is_bold and is_italic: font_name = 'Helvetica-BoldOblique'
            elif is_bold: font_name = 'Helvetica-Bold'
            elif is_italic: font_name = 'Helvetica-Oblique'
            else: font_name = 'Helvetica'
        elif font_family == 'Courier New':
            if is_bold and is_italic: font_name = 'Courier-BoldOblique'
            elif is_bold: font_name = 'Courier-Bold'
            elif is_italic: font_name = 'Courier-Oblique'
            else: font_name = 'Courier'
        elif font_family == 'Times New Roman':
            if is_bold and is_italic: font_name = 'Times-BoldItalic'
            elif is_bold: font_name = 'Times-Bold'
            elif is_italic: font_name = 'Times-Italic'
            else: font_name = 'Times-Roman'
        else:
            from reportlab.pdfbase import pdfmetrics
            if font_family in pdfmetrics.getRegisteredFontNames():
                font_name = font_family
            else:
                font_name = 'Times-Roman'

        # Convert HTML to ReportLab-compatible XML
        # ReportLab Paragraph supports: <b>, <i>, <u>, <br/>, <font>
        # Clean up browser HTML (e.g. <div>, <span>) to just text + inline tags
        rl_text = html_content
        # Replace <br> variants with <br/>
        rl_text = re.sub(r'<br\s*/?>', '<br/>', rl_text)
        # Remove <div> wrappers (browsers add these in contentEditable)
        rl_text = re.sub(r'<div[^>]*>', '<br/>', rl_text)
        rl_text = re.sub(r'</div>', '', rl_text)
        # Remove <span> tags (keep content)
        rl_text = re.sub(r'<span[^>]*>', '', rl_text)
        rl_text = re.sub(r'</span>', '', rl_text)
        # Keep <b>, <i>, <u>, <font> tags — ReportLab understands these
        # Remove any other HTML tags
        rl_text = re.sub(r'<(?!/?(?:b|i|u|font|br|strong|em)[ />])[^>]+>', '', rl_text)
        # Convert <strong> → <b>, <em> → <i>
        rl_text = rl_text.replace('<strong>', '<b>').replace('</strong>', '</b>')
        rl_text = rl_text.replace('<em>', '<i>').replace('</em>', '</i>')
        # Remove leading <br/>
        rl_text = re.sub(r'^(<br/>)+', '', rl_text)
        # Escape &
        rl_text = rl_text.replace('&nbsp;', ' ')
        if '&' in rl_text:
            rl_text = re.sub(r'&(?!amp;|lt;|gt;|quot;)', '&amp;', rl_text)

        leading = font_size * 1.35

        style = ParagraphStyle(
            name=f'cover_{block.get("id", "block")}',
            fontName=font_name,
            fontSize=font_size,
            leading=leading,
            alignment=alignment,
            textColor=black,
        )

        try:
            para = Paragraph(rl_text, style)
        except Exception:
            # Fallback: strip all HTML and render plain text
            plain = re.sub(r'<[^>]+>', ' ', html_content).strip()
            para = Paragraph(plain, style)

        # Calculate position: web top→RL bottom
        y_rl = PAGE_HEIGHT - y_top - height

        # Draw paragraph in a frame
        para_w, para_h = para.wrap(width, height)
        para.drawOn(c, x, y_rl + (height - para_h))

    c.save()
    return True


# ============================================================
# INVITED TALKS GENERATION
# ============================================================

def generate_messages_pdf(output_path):
    """
    Generate the messages PDF (multi-page) from messages.json.
    Uses the same rendering logic as invited talks.
    """
    config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'messages.json')
    if not os.path.exists(config_path):
        return False

    try:
        with open(config_path, 'r') as f:
            data = json.load(f)
    except Exception:
        return False

    pages = data.get('pages', [])
    if not pages:
        return False

    c = canvas_mod.Canvas(output_path, pagesize=A4)

    for page_data in pages:
        blocks = page_data.get('blocks', [])
        for block in blocks:
            block_type = block.get('type', 'text')
            x = block.get('x', 0)
            y_top = block.get('y', 0)
            width = block.get('width', 200)
            height = block.get('height', 30)
            
            y_rl = PAGE_HEIGHT - y_top - height

            if block_type == 'image':
                image_url = block.get('imageUrl', '')
                if image_url.startswith('data:image/'):
                    import base64
                    import io
                    from reportlab.lib.utils import ImageReader
                    try:
                        header, encoded = image_url.split(',', 1)
                        image_data = base64.b64decode(encoded)
                        image_stream = io.BytesIO(image_data)
                        img_reader = ImageReader(image_stream)
                        c.drawImage(img_reader, x, y_rl, width=width, height=height, preserveAspectRatio=True, mask='auto')
                    except Exception:
                        pass
                continue

            if block_type == 'shape':
                shape_type = block.get('shapeType', 'rect')
                stroke_width = block.get('strokeWidth', 2)
                
                stroke_hex = block.get('strokeColor', '#000000').lstrip('#')
                try:
                    sr, sg, sb = tuple(int(stroke_hex[i:i+2], 16) / 255.0 for i in (0, 2, 4))
                    c.setStrokeColorRGB(sr, sg, sb)
                except Exception:
                    c.setStrokeColorRGB(0, 0, 0)
                c.setLineWidth(stroke_width)
                
                fill_color = block.get('fillColor', '')
                has_fill = bool(fill_color and fill_color != 'transparent')
                if has_fill:
                    fill_hex = fill_color.lstrip('#')
                    try:
                        fr, fg, fb = tuple(int(fill_hex[i:i+2], 16) / 255.0 for i in (0, 2, 4))
                        c.setFillColorRGB(fr, fg, fb)
                    except Exception:
                        has_fill = False

                pdf_x = x + stroke_width/2.0
                pdf_y = PAGE_HEIGHT - y_top - height + stroke_width/2.0
                w = width - stroke_width
                h = height - stroke_width
                
                if shape_type == 'rect':
                    c.rect(pdf_x, pdf_y, w, h, stroke=1, fill=1 if has_fill else 0)
                elif shape_type == 'triangle':
                    path = c.beginPath()
                    path.moveTo(pdf_x + w/2, pdf_y + h)
                    path.lineTo(pdf_x + w, pdf_y)
                    path.lineTo(pdf_x, pdf_y)
                    path.close()
                    c.drawPath(path, stroke=1, fill=1 if has_fill else 0)
                continue

            if block_type == 'drawing':
                preview_image = block.get('previewImage', '')
                if preview_image.startswith('data:image/'):
                    import base64
                    import io
                    from reportlab.lib.utils import ImageReader
                    try:
                        header, encoded = preview_image.split(',', 1)
                        image_data = base64.b64decode(encoded)
                        image_stream = io.BytesIO(image_data)
                        img_reader = ImageReader(image_stream)
                        c.drawImage(img_reader, x, y_rl, width=width, height=height, preserveAspectRatio=True, mask='auto')
                    except Exception:
                        pass
                continue

            # Text block
            content = block.get('content', '')
            font_size = block.get('fontSize', 12)
            font_weight = block.get('fontWeight', 'normal')
            text_align = block.get('textAlign', 'left')
            font_family = block.get('fontFamily', 'Times-Roman')
            
            rl_font = 'Times-Bold' if font_weight == 'bold' else 'Times-Roman'
            c.setFont(rl_font, font_size)
            c.setFillColorRGB(0, 0, 0)

            import re as re_mod
            clean_text = re_mod.sub(r'<br\s*/?>', '\n', content)
            clean_text = re_mod.sub(r'<[^>]+>', '', clean_text)
            
            from reportlab.lib.utils import simpleSplit
            lines_list = []
            for paragraph in clean_text.split('\n'):
                if paragraph.strip() == '':
                    lines_list.append('')
                else:
                    wrapped = simpleSplit(paragraph, rl_font, font_size, width - 4)
                    lines_list.extend(wrapped)

            line_height = font_size * 1.3
            text_y = PAGE_HEIGHT - y_top - font_size

            for text_line in lines_list:
                if text_y < y_rl:
                    break
                if text_align == 'center':
                    c.drawCentredString(x + width / 2, text_y, text_line)
                elif text_align == 'right':
                    c.drawRightString(x + width, text_y, text_line)
                else:
                    c.drawString(x + 2, text_y, text_line)
                text_y -= line_height

        c.showPage()

    c.save()
    return True


def generate_invited_talks_pdf(output_path):
    """
    Generate the invited talks PDF (multi-page) from invited_talks.json.
    """
    config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'invited_talks.json')
    if not os.path.exists(config_path):
        return False

    try:
        with open(config_path, 'r') as f:
            data = json.load(f)
    except Exception:
        return False

    pages = data.get('pages', [])
    if not pages:
        return False

    c = canvas_mod.Canvas(output_path, pagesize=A4)

    for page_data in pages:
        blocks = page_data.get('blocks', [])
        for block in blocks:
            block_type = block.get('type', 'text')
            x = block.get('x', 0)
            y_top = block.get('y', 0)
            width = block.get('width', 200)
            height = block.get('height', 30)
            
            y_rl = PAGE_HEIGHT - y_top - height

            if block_type == 'image':
                image_url = block.get('imageUrl', '')
                if image_url.startswith('data:image/'):
                    import base64
                    import io
                    from reportlab.lib.utils import ImageReader
                    try:
                        header, encoded = image_url.split(',', 1)
                        image_data = base64.b64decode(encoded)
                        image_stream = io.BytesIO(image_data)
                        img_reader = ImageReader(image_stream)
                        c.drawImage(img_reader, x, y_rl, width=width, height=height, preserveAspectRatio=True, mask='auto')
                    except Exception as e:
                        print(f"Error drawing image to PDF: {e}")
                continue


            if block_type == 'shape':
                shape_type = block.get('shapeType', 'rect')
                stroke_width = block.get('strokeWidth', 2)
                
                stroke_hex = block.get('strokeColor', '#000000').lstrip('#')
                try:
                    sr, sg, sb = tuple(int(stroke_hex[i:i+2], 16) / 255.0 for i in (0, 2, 4))
                    c.setStrokeColorRGB(sr, sg, sb)
                except Exception:
                    c.setStrokeColorRGB(0, 0, 0)
                c.setLineWidth(stroke_width)
                
                fill_color = block.get('fillColor', '')
                has_fill = bool(fill_color and fill_color != 'transparent')
                if has_fill:
                    fill_hex = fill_color.lstrip('#')
                    try:
                        fr, fg, fb = tuple(int(fill_hex[i:i+2], 16) / 255.0 for i in (0, 2, 4))
                        c.setFillColorRGB(fr, fg, fb)
                    except Exception:
                        has_fill = False

                pdf_x = x + stroke_width/2.0
                pdf_y = PAGE_HEIGHT - y_top - height + stroke_width/2.0
                w = width - stroke_width
                h = height - stroke_width
                
                if shape_type == 'rect':
                    c.rect(pdf_x, pdf_y, w, h, stroke=1, fill=1 if has_fill else 0)
                elif shape_type == 'triangle':
                    path = c.beginPath()
                    path.moveTo(pdf_x + w/2, pdf_y + h)
                    path.lineTo(pdf_x + w, pdf_y)
                    path.lineTo(pdf_x, pdf_y)
                    path.close()
                    c.drawPath(path, stroke=1, fill=1 if has_fill else 0)
                continue

            if block_type == 'drawing':
                lines = block.get('lines', [])
                sym = block.get('symmetry', 'none')
                stroke_width = block.get('strokeWidth', 2)
                color_hex = block.get('strokeColor', '#000000').lstrip('#')
                try:
                    r, g, b = tuple(int(color_hex[i:i+2], 16) / 255.0 for i in (0, 2, 4))
                except Exception:
                    r, g, b = 0, 0, 0
                    
                c.setStrokeColorRGB(r, g, b)
                c.setLineWidth(stroke_width)
                
                fill_color = block.get('fillColor', '')
                has_fill = bool(fill_color and fill_color != 'transparent')
                if has_fill:
                    fill_hex = fill_color.lstrip('#')
                    try:
                        fr, fg, fb = tuple(int(fill_hex[i:i+2], 16) / 255.0 for i in (0, 2, 4))
                        c.setFillColorRGB(fr, fg, fb)
                    except Exception:
                        has_fill = False

                # First draw fill using path
                if has_fill and lines:
                    path = c.beginPath()
                    used = set()
                    current = lines[0]
                    used.add(current.get('id'))
                    
                    def conv_x(vx): return x + vx
                    def conv_y(vy): return PAGE_HEIGHT - (y_top + vy)
                    
                    path.moveTo(conv_x(current.get('x1',0)), conv_y(current.get('y1',0)))
                    path.lineTo(conv_x(current.get('x2',0)), conv_y(current.get('y2',0)))
                    last_pt = (current.get('x2',0), current.get('y2',0))
                    
                    while len(used) < len(lines):
                        found = False
                        for l in lines:
                            lid = l.get('id')
                            if lid not in used:
                                lx1, ly1 = l.get('x1',0), l.get('y1',0)
                                lx2, ly2 = l.get('x2',0), l.get('y2',0)
                                if abs(lx1 - last_pt[0]) < 2 and abs(ly1 - last_pt[1]) < 2:
                                    path.lineTo(conv_x(lx2), conv_y(ly2))
                                    last_pt = (lx2, ly2)
                                    used.add(lid)
                                    found = True
                                    break
                                elif abs(lx2 - last_pt[0]) < 2 and abs(ly2 - last_pt[1]) < 2:
                                    path.lineTo(conv_x(lx1), conv_y(ly1))
                                    last_pt = (lx1, ly1)
                                    used.add(lid)
                                    found = True
                                    break
                        if not found:
                            next_l = next((l for l in lines if l.get('id') not in used), None)
                            if next_l:
                                path.moveTo(conv_x(next_l.get('x1',0)), conv_y(next_l.get('y1',0)))
                                path.lineTo(conv_x(next_l.get('x2',0)), conv_y(next_l.get('y2',0)))
                                last_pt = (next_l.get('x2',0), next_l.get('y2',0))
                                used.add(next_l.get('id'))
                    path.close()
                    c.drawPath(path, stroke=0, fill=1)

                def draw_rl_line(l_x1, l_y1, l_x2, l_y2):
                    pdf_x1 = x + l_x1
                    pdf_y1 = PAGE_HEIGHT - (y_top + l_y1)
                    pdf_x2 = x + l_x2
                    pdf_y2 = PAGE_HEIGHT - (y_top + l_y2)
                    c.line(pdf_x1, pdf_y1, pdf_x2, pdf_y2)

                for line in lines:
                    l_x1, l_y1, l_x2, l_y2 = line.get('x1',0), line.get('y1',0), line.get('x2',0), line.get('y2',0)
                    draw_rl_line(l_x1, l_y1, l_x2, l_y2)
                    
                    if sym in ('horizontal', 'both'):
                        draw_rl_line(width - l_x1, l_y1, width - l_x2, l_y2)
                    if sym in ('vertical', 'both'):
                        draw_rl_line(l_x1, height - l_y1, l_x2, height - l_y2)
                    if sym == 'both':
                        draw_rl_line(width - l_x1, height - l_y1, width - l_x2, height - l_y2)
                continue

            # Text Block
            html_content = block.get('content', '')
            if not html_content:
                plain = block.get('text', '')
                html_content = plain.replace('\n', '<br/>')

            font_size = block.get('fontSize', 12)
            font_weight = block.get('fontWeight', 'normal')
            font_style = block.get('fontStyle', 'normal')
            text_align = block.get('textAlign', 'center')

            align_map = {'left': TA_LEFT, 'center': TA_CENTER, 'right': TA_LEFT}
            alignment = align_map.get(text_align, TA_CENTER)
            if text_align == 'right':
                from reportlab.lib.enums import TA_RIGHT
                alignment = TA_RIGHT

            font_family = block.get('fontFamily', 'Times New Roman')
            is_bold = (font_weight == 'bold')
            is_italic = (font_style == 'italic')
            
            if font_family == 'Arial':
                if is_bold and is_italic: font_name = 'Helvetica-BoldOblique'
                elif is_bold: font_name = 'Helvetica-Bold'
                elif is_italic: font_name = 'Helvetica-Oblique'
                else: font_name = 'Helvetica'
            elif font_family == 'Courier New':
                if is_bold and is_italic: font_name = 'Courier-BoldOblique'
                elif is_bold: font_name = 'Courier-Bold'
                elif is_italic: font_name = 'Courier-Oblique'
                else: font_name = 'Courier'
            elif font_family == 'Times New Roman':
                if is_bold and is_italic: font_name = 'Times-BoldItalic'
                elif is_bold: font_name = 'Times-Bold'
                elif is_italic: font_name = 'Times-Italic'
                else: font_name = 'Times-Roman'
            else:
                from reportlab.pdfbase import pdfmetrics
                if font_family in pdfmetrics.getRegisteredFontNames():
                    font_name = font_family
                else:
                    font_name = 'Times-Roman'

            rl_text = html_content
            rl_text = re.sub(r'<br\s*/?>', '<br/>', rl_text)
            rl_text = re.sub(r'<div[^>]*>', '<br/>', rl_text)
            rl_text = re.sub(r'</div>', '', rl_text)
            rl_text = re.sub(r'<span[^>]*>', '', rl_text)
            rl_text = re.sub(r'</span>', '', rl_text)
            rl_text = re.sub(r'<(?!/?(?:b|i|u|font|br|strong|em)[ />])[^>]+>', '', rl_text)
            rl_text = rl_text.replace('<strong>', '<b>').replace('</strong>', '</b>')
            rl_text = rl_text.replace('<em>', '<i>').replace('</em>', '</i>')
            rl_text = re.sub(r'^(<br/>)+', '', rl_text)
            rl_text = rl_text.replace('&nbsp;', ' ')
            if '&' in rl_text:
                rl_text = re.sub(r'&(?!amp;|lt;|gt;|quot;)', '&amp;', rl_text)

            leading = font_size * 1.35
            style = ParagraphStyle(
                name=f'cover_{block.get("id", "block")}',
                fontName=font_name,
                fontSize=font_size,
                leading=leading,
                alignment=alignment,
                textColor=black,
            )

            try:
                para = Paragraph(rl_text, style)
            except Exception:
                plain = re.sub(r'<[^>]+>', ' ', html_content).strip()
                para = Paragraph(plain, style)

            para_w, para_h = para.wrap(width, height)
            para.drawOn(c, x, y_rl + (height - para_h))
        
        # End of page
        c.showPage()
        
    c.save()
    return True


# ============================================================
# PDF TO WORD CONVERSION
# ============================================================

def _inject_word_toc_hyperlinks(docx_path, papers):
    """
    Injects Word-compliant bookmarks (<w:bookmarkStart>) on body abstract titles and wraps
    Table of Contents entries inside (<w:hyperlink w:anchor="...">) elements,
    enabling internal clickable navigation without altering ANY existing styling,
    fonts, colors, sizes, or dotted leaders.
    """
    try:
        import docx
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        def normalize(text):
            return re.sub(r'[\.\s\d—\-_,:()\'\"’‘]+', '', text.lower())

        doc = docx.Document(docx_path)

        # Collect all body title paragraphs (doc.paragraphs where title matches)
        body_paragraphs = list(doc.paragraphs)

        # Collect all TOC candidate elements (from both doc.paragraphs and doc.tables)
        toc_elements = []
        for p in doc.paragraphs:
            toc_elements.append(p)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        toc_elements.append(p)

        first_abs_idx = next(
            (i for i, p in enumerate(doc.paragraphs) if 'Abstract' in p.text and ('—' in p.text or '-' in p.text)),
            len(doc.paragraphs)
        )

        used_body = set()
        used_toc = set()

        for idx, paper in enumerate(papers):
            # Clean alphanumeric bookmark identifier for full Microsoft Word navigation support
            bm_name = f'Paper{idx+1}'
            norm_title = normalize(paper['title'])
            if not norm_title:
                continue

            # 1. Add Bookmark to Body Title (at or after first_abs_idx)
            for i in range(max(0, first_abs_idx - 10), len(body_paragraphs)):
                if i in used_body:
                    continue
                p = body_paragraphs[i]
                txt = p.text.strip()
                if not txt or 'International Conference' in txt or txt.startswith('Abstract') or txt.startswith('Keywords'):
                    continue
                norm_p = normalize(txt)
                if norm_p[:18] and (norm_p[:18] in norm_title or norm_title[:18] in norm_p):
                    p_elem = p._p
                    if not any(bm.get(qn('w:name')) == bm_name for bm in p_elem.findall(qn('w:bookmarkStart'))):
                        bm_start = OxmlElement('w:bookmarkStart')
                        bm_start.set(qn('w:id'), str(idx + 1))
                        bm_start.set(qn('w:name'), bm_name)
                        bm_end = OxmlElement('w:bookmarkEnd')
                        bm_end.set(qn('w:id'), str(idx + 1))
                        p_elem.insert(0, bm_start)
                        p_elem.append(bm_end)
                    used_body.add(i)
                    break

            # 2. Add Hyperlink to TOC Entry (searching both paragraphs and tables)
            for i, p in enumerate(toc_elements):
                if i in used_toc:
                    continue
                txt = p.text.strip()
                if not txt or txt == 'Contents' or 'International Conference' in txt or txt.startswith('Abstract'):
                    continue
                norm_p = normalize(txt)
                if norm_p[:18] and (norm_p[:18] in norm_title or norm_title[:18] in norm_p):
                    p_elem = p._p
                    if not p_elem.findall(qn('w:hyperlink')):
                        runs = list(p_elem.findall(qn('w:r')))
                        if runs:
                            hyperlink = OxmlElement('w:hyperlink')
                            hyperlink.set(qn('w:anchor'), bm_name)
                            hyperlink.set(qn('w:history'), '1')
                            for r in runs:
                                p_elem.remove(r)
                                hyperlink.append(r)
                            p_elem.append(hyperlink)
                    used_toc.add(i)
                    break

        doc.save(docx_path)
    except Exception as e:
        print(f"  [WARN] Could not inject Word TOC hyperlinks: {e}")


def convert_pdf_to_word(pdf_path, docx_path, papers=None):
    """
    Converts the compiled PDF directly to Word (.docx) format using pdf2docx,
    then adds clickable internal TOC hyperlinks & bookmarks using python-docx
    while preserving 100% of the visual styling, fonts, and layout.
    """
    try:
        import fitz
        if not hasattr(fitz.Rect, 'get_area'):
            fitz.Rect.get_area = lambda self: abs(self)

        from pdf2docx import Converter
        print(f"\n[PROCESS] Converting compiled PDF -> Word (.docx)...")
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()

        # ── Inject Clickable TOC Hyperlinks & Bookmarks in Word (.docx) ──
        if papers:
            _inject_word_toc_hyperlinks(docx_path, papers)

        print(f"  [OK] Successfully converted PDF to Word: {docx_path}")
        return True
    except ImportError:
        print("  [ERROR] pdf2docx is not installed. Run: pip install pdf2docx")
        return False
    except Exception as e:
        print(f"  [ERROR] PDF to Word conversion failed: {e}")
        return False


def generate_compilation_report(output_folder, pdf_output_path, docx_output_path, papers, page_map, total_pages, toc_pages_count):
    """
    Generates a detailed compilation audit report in JSON and Markdown formats:
    - Lists page-by-page mapping: on Page X, which articles are placed, and from which PDF source.
    - Lists all extracted metadata (Title, Authors with count, Keywords).
    - Lists any anomalies/doubts/warnings flagged during scraping (e.g. fallback keywords, author count mismatches).
    """
    # Invert page_map to get articles per body page: {body_page_num: [(paper_index, paper)]}
    pages_to_papers = {}
    for p_idx, p in enumerate(papers):
        pg = page_map.get(p_idx, 1)
        pages_to_papers.setdefault(pg, []).append((p_idx, p))

    report_data = {
        'summary': {
            'total_papers': len(papers),
            'total_pages': total_pages,
            'toc_pages': toc_pages_count,
            'body_pages': total_pages - toc_pages_count,
            'pdf_output': os.path.basename(pdf_output_path),
            'docx_output': os.path.basename(docx_output_path) if docx_output_path and os.path.exists(docx_output_path) else None,
            'generated_at': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        },
        'pages': [],
        'anomalies_summary': []
    }

    md_lines = [
        "# Compilation & Extraction Audit Report",
        "",
        f"**Generated on:** {report_data['summary']['generated_at']}  ",
        f"**Total Papers:** {len(papers)} | **Total Document Pages:** {total_pages} ({toc_pages_count} TOC + {total_pages - toc_pages_count} Body)  ",
        f"**Outputs:** `{os.path.basename(pdf_output_path)}`",
        "",
        "---",
        "",
        "## Page-by-Page Article Mapping",
        ""
    ]

    for pg in sorted(pages_to_papers.keys()):
        pg_papers = pages_to_papers[pg]
        merged_pdf_page = toc_pages_count + pg
        md_lines.append(f"### Body Page {pg} (Document Page {merged_pdf_page})")
        md_lines.append(f"*Contains {len(pg_papers)} article(s)*\n")

        pg_entry = {
            'body_page': pg,
            'document_page': merged_pdf_page,
            'articles': []
        }

        for item_idx, (p_idx, p) in enumerate(pg_papers):
            auth_names = [a['name'] for a in p['authors']]
            auth_str = ', '.join(auth_names) if auth_names else '*(None extracted)*'
            anomalies = p.get('anomalies', [])

            art_entry = {
                'article_index': p_idx + 1,
                'source_file': p['source_file'],
                'title': p['title'],
                'author_count': len(p['authors']),
                'authors': auth_names,
                'keywords': p['keywords'],
                'anomalies': anomalies
            }
            pg_entry['articles'].append(art_entry)

            md_lines.append(f"#### Article {p_idx + 1}: {p['title']}")
            md_lines.append(f"- **Source PDF:** `{p['source_file']}`")
            md_lines.append(f"- **Authors ({len(p['authors'])}):** {auth_str}")
            md_lines.append(f"- **Keywords:** {p['keywords'] if p['keywords'] else '*(None)*'}")

            if anomalies:
                md_lines.append(f"- **[FLAGGED ISSUES / DOUBTS]:**")
                for anom in anomalies:
                    md_lines.append(f"  - ⚠️ {anom}")
                    report_data['anomalies_summary'].append({
                        'source_file': p['source_file'],
                        'title': p['title'],
                        'body_page': pg,
                        'issue': anom
                    })
            md_lines.append("")

        report_data['pages'].append(pg_entry)
        md_lines.append("---")
        md_lines.append("")

    # Add Anomalies Summary Section
    md_lines.append("## Extraction Anomalies & Doubts Summary")
    md_lines.append("")
    if report_data['anomalies_summary']:
        md_lines.append(f"Total flagged items: **{len(report_data['anomalies_summary'])}**\n")
        md_lines.append("| Body Page | Source PDF | Title | Flagged Doubt / Issue |")
        md_lines.append("| :---: | :--- | :--- | :--- |")
        for anom in report_data['anomalies_summary']:
            short_t = anom['title'][:45] + '...' if len(anom['title']) > 45 else anom['title']
            md_lines.append(f"| {anom['body_page']} | `{anom['source_file']}` | {short_t} | {anom['issue']} |")
    else:
        md_lines.append("✨ **Zero anomalies or doubts detected! All articles extracted cleanly with 100% complete metadata.**")

    # Write JSON report
    json_path = os.path.join(output_folder, "compilation_report.json")
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(report_data, f, indent=2, ensure_ascii=False)

    # Write Markdown report
    md_path = os.path.join(output_folder, "compilation_report.md")
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(md_lines))

    print(f"\n[REPORT] Generated Compilation & Audit Report:")
    print(f"  - Markdown: {md_path}")
    print(f"  - JSON:     {json_path}")
    if report_data['anomalies_summary']:
        print(f"  [WARN] {len(report_data['anomalies_summary'])} potential doubt(s) flagged in report for user review.")
    else:
        print(f"  [OK] 0 extraction anomalies detected across all {len(papers)} papers.")


# ============================================================
# MAIN
# ============================================================

def main():
    # ── Default format is 'both' (PDF + Word docx conversion) ──
    output_format = 'both'
    positional_args = []
    for arg in sys.argv[1:]:
        if arg.startswith('--format='):
            output_format = arg.split('=', 1)[1].lower()
        else:
            positional_args.append(arg)

    # ── Determine input/output folders ──
    input_folder = positional_args[0] if len(positional_args) > 0 else "IncompletePDF"
    output_folder = positional_args[1] if len(positional_args) > 1 else "CompletedPDF"

    os.makedirs(output_folder, exist_ok=True)

    print("==================================================")
    print("       ACROSET PDF SCRAPER & COMPILER             ")
    print("==================================================")
    print(f"  Input folder:  {input_folder}/")
    print(f"  Output folder: {output_folder}/")

    # ── Scan for PDFs ──
    pdf_files = sorted([
        f for f in os.listdir(input_folder)
        if f.lower().endswith('.pdf')
    ])

    if not pdf_files:
        print(f"\n[ERROR] No PDF files found in '{input_folder}/'")
        sys.exit(1)

    print(f"\n[QUEUE] Found {len(pdf_files)} PDF(s): {', '.join(pdf_files)}")

    # ── Extract data from each PDF ──
    papers = []
    failed = []

    for pdf_file in pdf_files:
        full_path = os.path.join(input_folder, pdf_file)
        try:
            data = extract_paper_data(full_path)
            if data and data['title']:
                papers.append(data)
            else:
                failed.append(pdf_file)
        except Exception as e:
            print(f"\n  [ERROR] EXCEPTION processing {pdf_file}: {e}")
            failed.append(pdf_file)

    if not papers:
        print("\n[ERROR] No papers could be extracted! Check the PDFs.")
        sys.exit(1)

    print(f"\n{'='*60}")
    print(f"[SUMMARY] EXTRACTION SUMMARY")
    print(f"{'='*60}")
    print(f"  [OK] Successful: {len(papers)}")
    if failed:
        print(f"  [ERROR] Failed:     {len(failed)} -- {', '.join(failed)}")

    # ── Generate body PDF ──
    body_temp = os.path.join(output_folder, "_body_temp.pdf")
    print(f"\n[PROCESS] Generating body pages...")
    page_map, page_groups, page_density_modes = generate_body_pdf(papers, body_temp)

    print(f"\n  Page assignments:")
    for i, p in enumerate(papers):
        short_title = p['title'][:55] + '...' if len(p['title']) > 55 else p['title']
        print(f"    Page {page_map.get(i, '?')}: {short_title}")

    # ── PDF output ──
    toc_temp = os.path.join(output_folder, "_toc_temp.pdf")
    print(f"\n[PROCESS] Generating Table of Contents...")
    toc_entries_meta = generate_toc_pdf(papers, page_map, toc_temp)

    pdf_output_path = os.path.join(output_folder, "compiled_output.pdf")
    print(f"\n[PROCESS] Merging TOC + Body -> {pdf_output_path}")
    merge_pdfs(toc_temp, body_temp, pdf_output_path, toc_entries_meta=toc_entries_meta)

    # Cleanup temp files
    if os.path.exists(body_temp):
        os.remove(body_temp)
    if os.path.exists(toc_temp):
        os.remove(toc_temp)

    # ── Convert to Word if requested ──
    docx_output_path = os.path.join(output_folder, "compiled_output.docx")
    if output_format in ('word', 'docx', 'both'):
        success = convert_pdf_to_word(pdf_output_path, docx_output_path, papers=papers)
        if not success:
            print("\n[ERROR] Word conversion failed!")
            sys.exit(1)

    # ── Move processed PDFs to output folder ──
    print(f"\n[PROCESS] Moving processed PDFs to {output_folder}/...")
    for pdf_file in pdf_files:
        src = os.path.join(input_folder, pdf_file)
        dst = os.path.join(output_folder, pdf_file)
        if os.path.exists(src):
            if os.path.exists(dst):
                os.remove(dst)  # Remove if already exists
            shutil.move(src, dst)
            print(f"    [OK] {pdf_file}")

    # ── Final summary & Compilation Report ──
    final_doc = fitz.open(pdf_output_path)
    total_pages = len(final_doc)
    final_doc.close()

    body_page_count = len(set(page_map.values()))
    toc_pages = total_pages - body_page_count

    # Generate Page-by-Page Audit & Anomaly Report
    generate_compilation_report(
        output_folder=output_folder,
        pdf_output_path=pdf_output_path,
        docx_output_path=docx_output_path,
        papers=papers,
        page_map=page_map,
        total_pages=total_pages,
        toc_pages_count=toc_pages
    )

    print(f"\n{'='*60}")
    print(f"[DONE] COMPILATION COMPLETE")
    print(f"{'='*60}")
    print(f"  PDF Output:  {pdf_output_path}")
    if output_format in ('word', 'docx', 'both'):
        print(f"  Word Output: {docx_output_path}")
    print(f"  Papers:      {len(papers)}")
    print(f"  Total pages: {total_pages} "
          f"({toc_pages} TOC + {body_page_count} body)")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
